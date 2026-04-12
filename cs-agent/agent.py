import json
import os
import time
from datetime import datetime
from openai import OpenAI
from config import CONFIG

# Singleton client — reuse instead of creating per request
_client = None


def get_client():
    global _client
    if _client is None:
        api_key = CONFIG["api_key"] or "lm-studio"  # No API key needed for local model
        _client = OpenAI(api_key=api_key, base_url=CONFIG["api_base"])
    return _client

# =============================================================================
# SYSTEM PROMPT — Sections 1, 2, 3, 10, 14 + Self-Check Rule
# =============================================================================

SYSTEM_PROMPT = """You are the AI assistant of the firm's Customer Success department. Your name is **CS Agent**.

Your mission: Support the coordination of accounting, payroll, tax compliance, immigration processes, and administrative procedures for client companies operating in Poland. You are used as an internal operational tool by the support team (Dogukan, Rana, Mehtap); you also prepare message drafts to be sent to clients.

---

## 1. CORE BEHAVIORAL RULES

### 1.1 Flow Rule (HIGHEST PRIORITY RULE)
**WORK IN A SINGLE STEP:** Whatever the user requests, call the necessary tools and return the result as a SINGLE RESPONSE.
- NEVER ask "Do you approve?" — neither for read nor write operations.
- Do NOT write intermediate summaries, plans, or step lists — deliver the result directly.
- If the user says "write a message": read the wiki + write the message + present it in a single response. Do not ask for approval in between.
- If the user says "yes" or "I approve": execute the action directly without producing another summary.
- Do NOT use the "I understood" format — this format has been removed.
- Read wiki, query mastersheet, prepare message — do ALL in a SINGLE STEP and show the result.
- The mastersheet_read tool supports general queries: "rastgele" (10 random firms), "tumu" (full list), "say" (total count), firm type ("JDG"), accountant name ("Kasia"). Query whatever the user requests directly, do not say "I cannot".

### 1.3 Output Format Rule (PRIORITY)
Write your responses in NATURAL LANGUAGE. Do NOT use markdown formatting:
- Do NOT use markdown markers like "---", "###", "**bold**", "- bullet".
- When providing information, use plain text paragraphs; list items within sentences instead of using bullet lists.
- When preparing message drafts, write them as REAL MESSAGES — natural text ready to send, not markdown.
- Only the [SOURCE:] tag may remain as markdown.

### 1.4 Deterministic Process Tracking
- Follow each process step by step. Do not make assumptions.
- If you are unsure about something, state it clearly and suggest consulting the relevant team member.
- Always clearly specify critical information such as dates, NIP, and company names.

### 1.5 Communication Language and Tone
- **Internal team** (Dogukan, Rana, Mehtap): Turkish, semi-formal, solution-oriented.
- **Accounting/Legal team** (Kasia, Liudmila, Aleksandra, Gosia, Jakub): English.
- **Client message drafts**: Turkish (if the client does not have a Turkish name, ask the user about the language). Semi-formal tone — friendly but professional. Avoid being overly casual.
- For technical/legal terms, provide both Turkish and Polish equivalents in parentheses: e.g., "VAT registration (rejestracja VAT)", "payroll (lista plac)".
- Message drafts are signed with "Support Team".

---

## 2. ORGANIZATION STRUCTURE

### 2.1 Core Team

| Person | Role | Area of Responsibility | Communication Language |
|--------|------|------------------------|------------------------|
| **Kaan Cakar** | General Manager / Financial Advisor | Strategic decisions, client escalations, financial advisory. Native Polish and Turkish speaker. Final decision-maker on many topics. | TR / PL |
| **Dogukan** | Customer Success Manager | Technical coordination, workflow tracking, operational process management | TR / EN |
| **Rana** | CS Specialist | Direct client communication, document collection, process tracking | TR |
| **Mehtap** | CS Specialist (new) | Direct client communication, document collection, process tracking | TR |
| **Kasia (Katarzyna)** | Head Accountant | Responsible for ~52 firms' accounting. Updates the mastersheet. Tax return preparation and submission. | PL / EN |
| **Liudmila** | Accountant | Responsible for ~30 firms' accounting | EN |
| **Aleksandra** | Document Specialist | Document control, filing, archiving | PL / EN |
| **Gosia** | Payroll Specialist | ZUS declarations, payroll calculations, A1 processes | PL / EN |
| **Jakub** | Lawyer | S24 company formation, name/ownership changes, company suspension. Active. | PL / EN |
| Karol | Lawyer (former) | **No longer active.** Do not route to him. | — |

### 2.2 Partner Firms

| Partner | Contact Person | Number of Firms | Relationship |
|---------|---------------|-----------------|--------------|
| **Smyrna** | Koray Binler | ~12 firms | Document collection and client communication on partner side. VAT registration, accounting, and official processes on the firm's side. |
| **Nesfa** | Ayse (former employee) | ~13 firms | Same model. |
| **G&L** | Gediz | ~54 firms | Same model. |

### 2.3 Routing Matrix

| Topic | Route To |
|-------|----------|
| Payroll (bordro), ZUS, salary calculation, A1 | **Gosia** |
| Accounting (muhasebe), tax returns, invoices, JPK | **Kasia** or **Liudmila** (based on the accountant assigned to the firm) |
| Missing documents, filing | **Aleksandra** |
| Company formation (S24), name/ownership change, structural legal matters | **Jakub** |
| Strategic decisions, major client escalation, financial advisory | **Kaan Bey** |
| Client communication, document follow-up, reminders | **Rana / Mehtap / Dogukan** |
| Payment/invoice reminders | **Get approval from support team first**, then forward to client |
| Contract cancellation request | **Write a report to Kaan Bey** + prepare a "We have received your request, we will get back to you shortly" message to the client + notify the support team |

---

## 3. SERVICE SCOPE

### 3.1 Accounting Services (Uslugi ksiegowe)
- Bookkeeping (pelna ksiegowosc / KPiR)
- VAT reconciliation and JPK_V7 filing
- CIT/PIT advance calculation and annual returns
- Financial reports (balance sheet, income statement)
- Tax return submission to the tax office

### 3.2 Payroll / HR Services (Kadry i place)
- Employment contract preparation (umowa o prace, zlecenie, dzielo)
- Salary calculation (gross-to-net) and payroll lists
- ZUS declarations (ZUA/ZZA/ZWUA/DRA/RCA)
- PIT-11, PIT-4R, PIT-8AR preparation
- Onboarding/offboarding procedures
- Leave and absence tracking

### 3.3 Immigration Services (Wsparcie imigracyjne)
- Application preparation and document compilation
- Document verification (format, completeness)
- Institutional communication and correspondence coordination
- Process tracking and status updates
- Translation coordination

### 3.4 Information and Administrative Support
- General information, checklists, and routing
- Administrative document preparation (forms, power of attorney, etc.)
- HR/Compliance reminders and checklists
- Audit/investigation preparation support

### 3.5 OUT OF SCOPE
- Financial audit, independent audit, forensic
- Court representation, litigation petitions, appeals
- Tax advisory (only informational guidance is provided, no advice)
- M&A, international tax planning
- Legal review of contracts

---

## 10. MESSAGE DRAFT RULES

Message drafts are written as READY-TO-SEND, NATURAL TEXT. Do NOT use markdown markers (---, ###, **, - ).

### 10.1 WhatsApp Messages (Turkish)
Short, clear, action-oriented. Example format:

Sayin [Name],

Support Team olarak [topic] hakkinda bilgilendirme yapmak istiyoruz. [Content paragraphs]

Iyi calismalar,
Support Team

### 10.2 Email Messages (Turkish)
Subject line is specified separately. Written in paragraphs, sentences used instead of bullet points. Example:

Konu: [Company Name] - [Topic]

Sayin [Name],

[Introduction paragraph — purpose]

[Detail paragraph — required documents or information listed within sentences]

[Closing — next step]

Iyi calismalar,
Support Team

### 10.3 Internal Team Messages (English)
Short, structured. Company name + NIP are always specified. Written in natural English paragraphs.

---

## 14. CONSTRAINTS

1. When you prepare a message draft or document, show it to the user before sending — but do not ask for approval at intermediate steps.
2. Do not provide legal or tax advice — only provide informational guidance.
3. Do not speculate about uncertain tax rates or penalty amounts.
4. Do not unnecessarily repeat personal data (PESEL, full address, passport no.).
5. Do not share pricing information.
6. Do not comment on competitor firms.
7. Termination, service suspension, or initiating legal proceedings — only within Kaan Bey's authority.
8. Do not route to Karol — he is no longer active.

---

## PROACTIVE LEARNING RULE

When you learn new and valuable information during a conversation that is not in the wiki:
1. Notice the information (new procedure, updated date, corrected information)
2. Ask the user: "Would you like me to save this information to the wiki?"
3. If approved, save it to the APPROPRIATE FOLDER using wiki_write (vergi/, onboarding/, operasyon/, etc.)
4. If you need to update an existing article, first read it with wiki_read, then write the updated version with wiki_write.

This rule applies only to real, verifiable information. Do not write the user's personal opinions or temporary information to the wiki.

---

## SELF-CHECK RULE

At the end of every response, cite your source:
- If you used information from the wiki -> [SOURCE: article_name.md]
- For information not found in the wiki -> [SOURCE: general knowledge]
- If you used a draft/ source -> [SOURCE: draft/article_name.md — unverified]
- If you do not know -> "I cannot provide definitive information on this topic"
Do not include information in your response that you cannot cite a source for.
"""

# =============================================================================
# TOOLS — wiki_read + mastersheet_read + wiki_write + doc generation + exa + krs
# =============================================================================

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "wiki_read",
            "description": "Reads a wiki article's content. Choose the article path from the list in INDEX.md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "article_path": {
                        "type": "string",
                        "description": "Article path, e.g.: onboarding/checklist.md"
                    }
                },
                "required": ["article_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "mastersheet_read",
            "description": "Queries client information from the mastersheet. Searches by company name, NIP, firm type (JDG, SP), or accountant name. Special commands: 'rastgele' (10 random firms), 'tumu' (full firm list), 'say' (total firm count).",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Search term: company name, NIP number, firm type (JDG, SP, SC), accountant name, or any keyword"
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "wiki_write",
            "description": "Writes or updates a wiki article. Can write to any folder. For new articles, choose the appropriate folder (onboarding/, vergi/, operasyon/, bordro/, sozlesme/, iletisim/, referans/). To update an existing article, use the same path.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "Article path, e.g.: vergi/yeni_konu.md or onboarding/checklist.md (update)"
                    },
                    "content": {
                        "type": "string",
                        "description": "Article content (in markdown format)"
                    }
                },
                "required": ["filename", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_pdf",
            "description": "Creates a PDF document. IMPORTANT: Content should be RICH and DETAILED — not just headings, but explanatory paragraphs, bullet lists, and notes. Use UPPERCASE lines for headings (no digits or ':'), '- ' for lists. Write information in each section as plain text paragraphs. Do not produce short and empty PDFs.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Document title"},
                    "content": {"type": "string", "description": "Document content — plain text"},
                    "filename": {"type": "string", "description": "File name (without extension), e.g.: mine_food_report"},
                    "preview": {"type": "boolean", "description": "true=show draft without saving, false=save", "default": False}
                },
                "required": ["title", "content", "filename"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": "Creates a DOCX (Word) document. IMPORTANT: Content should be RICH and DETAILED — not just headings, but explanatory paragraphs, bullet lists, and notes. Use UPPERCASE lines for headings (no digits or ':'), '- ' for lists. Write information in each section as plain text paragraphs.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Document title"},
                    "content": {"type": "string", "description": "Document content — plain text"},
                    "filename": {"type": "string", "description": "File name (without extension), e.g.: mine_food_checklist"},
                    "preview": {"type": "boolean", "description": "true=show draft without saving, false=save", "default": False}
                },
                "required": ["title", "content", "filename"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_xlsx",
            "description": "Creates an Excel spreadsheet. Content is sent in CSV format: first row is headers, subsequent rows are data. Columns are separated by commas. File is saved to the outputs/ directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Sheet title"},
                    "content": {"type": "string", "description": "Table data in CSV format (first row is header)"},
                    "filename": {"type": "string", "description": "File name (without extension), e.g.: client_list"}
                },
                "required": ["title", "content", "filename"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "exa_search",
            "description": "Performs a web search. Use to find current information about accounting, tax, law, and Polish regulations. For general knowledge queries, procedure searches, and legislation checks.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Search query, e.g.: 'Poland VAT registration 2026 requirements'"
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "krs_lookup",
            "description": "Queries company information from the Polish KRS (Krajowy Rejestr Sadowy) API. Searches by KRS number. Returns company name, NIP, REGON, address, capital, and PKD code.",
            "parameters": {
                "type": "object",
                "properties": {
                    "krs_number": {
                        "type": "string",
                        "description": "KRS number, e.g.: 0001229771"
                    }
                },
                "required": ["krs_number"]
            }
        }
    }
]

# =============================================================================
# MCP CLIENT — Optional external tool integrations (Gmail, Supabase, etc.)
# =============================================================================

_mcp_processes = {}  # server_name -> subprocess.Popen
_mcp_tool_registry = {}  # tool_name -> server_name (built during discovery)
_mcp_request_id = 0


def _init_mcp_servers():
    """Start enabled MCP servers as subprocesses. Called once at agent startup."""
    import subprocess
    for name, cfg in CONFIG.get("mcp_servers", {}).items():
        if not cfg.get("enabled"):
            continue
        env = {**os.environ, **cfg.get("env", {})}
        try:
            proc = subprocess.Popen(
                cfg["command"],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=env,
            )
            _mcp_processes[name] = proc
            # Discover available tools from this server
            discovered = _discover_mcp_tools(name)
            for tool_def in discovered:
                tool_name = tool_def.get("name", "")
                if tool_name:
                    _mcp_tool_registry[tool_name] = name
            print(f"MCP server '{name}' started (pid={proc.pid}, {len(discovered)} tools)")
        except FileNotFoundError:
            print(f"MCP server '{name}' failed to start — command not found: {cfg['command'][0]}")
        except Exception as e:
            print(f"MCP server '{name}' failed to start — {e}")


def _mcp_request(server_name, method, params=None):
    """Send a JSON-RPC request to an MCP server and return the result."""
    global _mcp_request_id
    proc = _mcp_processes.get(server_name)
    if not proc or proc.poll() is not None:
        return {"error": f"MCP server '{server_name}' is not running"}

    _mcp_request_id += 1
    request = {
        "jsonrpc": "2.0",
        "id": _mcp_request_id,
        "method": method,
        "params": params or {},
    }
    try:
        request_line = json.dumps(request) + "\n"
        proc.stdin.write(request_line.encode())
        proc.stdin.flush()
        response_line = proc.stdout.readline().decode()
        return json.loads(response_line)
    except Exception as e:
        return {"error": f"MCP communication error: {e}"}


def _discover_mcp_tools(server_name):
    """Query an MCP server for its available tools. Returns list of tool definitions."""
    result = _mcp_request(server_name, "tools/list")
    if "error" in result:
        return []
    return result.get("result", {}).get("tools", [])


def _call_mcp_tool(server_name, tool_name, args):
    """Execute a tool call on an MCP server. Returns result string."""
    result = _mcp_request(server_name, "tools/call", {
        "name": tool_name,
        "arguments": args,
    })
    if "error" in result:
        return f"Error: MCP tool call failed — {result['error']}"
    content = result.get("result", {}).get("content", [])
    if content and isinstance(content, list):
        return "\n".join(c.get("text", str(c)) for c in content)
    return str(result.get("result", ""))


def _shutdown_mcp_servers():
    """Gracefully terminate all running MCP server subprocesses."""
    for name, proc in _mcp_processes.items():
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except Exception:
                proc.kill()
    _mcp_processes.clear()


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================


def read_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def safe_nip(raw):
    """NIP float->int conversion. Mastersheet stores NIP as float."""
    try:
        return str(int(float(raw)))
    except (ValueError, TypeError):
        return str(raw).strip()


def format_client_row(row):
    nip = safe_nip(row.get("NIP/PESEL", ""))
    return (
        f"Company: {row.get('Company Name', '?')}\n"
        f"NIP: {nip}\n"
        f"Type: {row.get('TYP', '?')}\n"
        f"KRS: {row.get('KRS', '?')}\n"
        f"Partner: {row.get('Responsible', '?')}\n"
        f"Accountant: {row.get('RA', '?')}\n"
        f"VAT: {row.get('VAT Aktiflik Tarihi', '?')}\n"
        f"Tax Office: {row.get('US', '?')}\n"
        f"Mikrorachunek: {row.get('Mikrorachunek', '?')}"
    )


# =============================================================================
# TOOL IMPLEMENTATIONS
# =============================================================================


def wiki_read(article_path: str) -> str:
    wiki_dir = os.path.realpath(CONFIG["wiki_dir"])
    full_path = os.path.realpath(f"{CONFIG['wiki_dir']}/{article_path}")
    if not full_path.startswith(wiki_dir):
        return f"Error: '{article_path}' invalid path — access outside wiki directory blocked."
    if not os.path.exists(full_path):
        return f"Error: '{article_path}' not found."
    content = read_file(full_path)
    if len(content) > CONFIG["wiki_truncate_at"] * 4:
        content = content[:CONFIG["wiki_truncate_at"] * 4] + "\n\n... [truncated]"
    if "draft/" in article_path:
        return "[WARNING: UNVERIFIED SOURCE]\n\n" + content
    return content


def mastersheet_read(query: str) -> str:
    import csv
    q = query.lower().strip()

    # Special commands: random pick, list all, count
    if q in ("rastgele", "random", "rastgele sec", "rastgele 10", "10 rastgele"):
        import random
        all_rows = []
        with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                if row.get("Company Name", "").strip():
                    all_rows.append(row)
        sample = random.sample(all_rows, min(10, len(all_rows)))
        header = f"(Random {len(sample)} firms / total {len(all_rows)})\n\n"
        return header + "\n---\n".join(format_client_row(r) for r in sample)

    if q in ("tumu", "hepsi", "tum firmalar", "liste", "listele", "all"):
        all_rows = []
        with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                name = row.get("Company Name", "").strip()
                if name:
                    all_rows.append(name)
        return f"Total {len(all_rows)} firms:\n\n" + "\n".join(f"- {n}" for n in all_rows[:50])

    if q in ("say", "kac firma", "toplam", "count"):
        with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
            count = sum(1 for row in csv.DictReader(f) if row.get("Company Name", "").strip())
        return f"Total {count} firms registered in the mastersheet."

    results = []
    partial = []
    with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            nip = safe_nip(row.get("NIP/PESEL", ""))
            company = row.get("Company Name", "")
            typ = row.get("TYP", "").strip()
            ra = row.get("RA", "").strip()
            cl = company.lower()
            tl = typ.lower()
            rl = ra.lower()

            # Exact match: company name, NIP, firm type, accountant
            if q in cl or q == nip or q == tl or q in rl:
                results.append(format_client_row(row))
            # Firm type match: abbreviations like "jdg", "sp", "sc"
            elif q in tl or tl.startswith(q):
                results.append(format_client_row(row))
            # Fuzzy: word-based matching (typo tolerance, min 3 chars in common)
            elif not results:
                q_words = [w for w in q.split() if len(w) >= 3]
                c_words = [w for w in cl.split() if len(w) >= 3]
                if q_words and c_words:
                    match_count = sum(1 for w in q_words if any(
                        w[:3] == cw[:3] and len(w) >= 3 for cw in c_words
                    ))
                    if match_count == len(q_words):
                        partial.append(format_client_row(row))
    if results:
        header = f"({len(results)} results found)\n\n" if len(results) > 1 else ""
        return header + "\n---\n".join(results[:10])
    if partial:
        return f"({len(partial)} partial matches)\n\n" + "\n---\n".join(partial[:5])
    return f"No matching firm found for '{query}'."


def wiki_write(filename: str, content: str) -> str:
    """Writes or updates a wiki article. T14: Automatic backup before writing."""
    if not filename.endswith(".md"):
        filename += ".md"
    wiki_dir = os.path.realpath(CONFIG["wiki_dir"])
    target = os.path.realpath(f"{CONFIG['wiki_dir']}/{filename}")
    if not target.startswith(wiki_dir):
        return f"Error: '{filename}' invalid — access outside wiki directory blocked."
    os.makedirs(os.path.dirname(target), exist_ok=True)
    is_update = os.path.exists(target)

    # T14: Version the existing file if it exists
    if is_update:
        _wiki_backup(filename, target)

    with open(target, "w", encoding="utf-8") as f:
        f.write(content)
    # Add to INDEX.md (if not present)
    index_path = CONFIG["index_file"]
    index_content = read_file(index_path)
    if filename not in index_content:
        index_content += f"\n- {filename}\n"
        with open(index_path, "w", encoding="utf-8") as f:
            f.write(index_content)
    action = "updated" if is_update else "created"
    return f"Wiki article {action}: {filename}"


def _wiki_backup(filename: str, filepath: str):
    """T14: Backup wiki article under .versions/. Max 10 versions."""
    import shutil
    versions_dir = os.path.join(CONFIG["wiki_dir"], ".versions", filename.replace("/", "_"))
    os.makedirs(versions_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = os.path.join(versions_dir, f"{ts}.md")
    shutil.copy2(filepath, backup_path)
    # Max 10 versions — oldest gets deleted
    versions = sorted(os.listdir(versions_dir))
    while len(versions) > 10:
        os.remove(os.path.join(versions_dir, versions.pop(0)))


# =============================================================================
# DOCUMENT GENERATION TOOLS
# =============================================================================


def _sanitize_filename(filename):
    """Sanitize filename for safe file system use."""
    import re
    return re.sub(r'[^\w\-]', '_', filename)


def _parse_content(content):
    """4-rule parsing: plain text, UPPERCASE heading, '- ' bullet, 'Key: Value' field.
    T15: field type added — key is rendered bold, value normal in PDF/DOCX.
    """
    import re
    lines = []
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            lines.append(("blank", ""))
        elif stripped.startswith("- "):
            lines.append(("bullet", stripped[2:]))
        elif ":" in stripped and not stripped.endswith(":"):
            # T15: Field detection — "Key: Value" format
            colon_pos = stripped.index(":")
            key_part = stripped[:colon_pos].strip()
            val_part = stripped[colon_pos+1:].strip()
            # Key part: short (max 30 char), may contain spaces, no digits (except dates)
            is_field = (
                2 <= len(key_part) <= 30
                and val_part
                and not key_part.startswith("http")
                and key_part[0].isalpha()
            )
            if is_field:
                lines.append(("field", f"{key_part}:{val_part}"))
            else:
                lines.append(("text", stripped))
        else:
            # Heading: letters + spaces only, at least 2 letters, all uppercase
            alpha_only = re.sub(r'[^a-zA-ZçÇğĞıİöÖşŞüÜ ]', '', stripped)
            is_heading = (
                len(alpha_only) >= 4
                and alpha_only == alpha_only.upper()
                and ":" not in stripped
                and not any(c.isdigit() for c in stripped)
                and stripped[0].isalpha()
            )
            lines.append(("heading" if is_heading else "text", stripped))
    return lines


def _content_summary(parsed_lines, doc_type):
    """Generate document summary."""
    headings = sum(1 for t, _ in parsed_lines if t == "heading")
    bullets = sum(1 for t, _ in parsed_lines if t == "bullet")
    texts = sum(1 for t, _ in parsed_lines if t == "text")
    return f"{doc_type.upper()} summary: {headings} headings, {bullets} bullets, {texts} paragraphs"


def _is_legal(title, content):
    """Legal document check — branding is not applied."""
    check = (title + " " + content[:200]).lower()
    return any(w in check for w in ["legal", "hukuki", "umowa", "aneks", "wypowiedzenie"])

def _fix_encoding(text):
    """Fix double-encoded UTF-8 (model sometimes sends broken encoding)."""
    try:
        return text.encode('latin-1').decode('utf-8')
    except (UnicodeDecodeError, UnicodeEncodeError):
        return text

def _create_pdf(title, content, filename, preview=False):
    from fpdf import FPDF

    title = _fix_encoding(title)
    content = _fix_encoding(content)

    # Brand — minimal palette
    ORANGE = (249, 115, 22)      # #F97316
    DARK = (31, 41, 55)          # #1F2937
    GRAY = (156, 163, 175)       # #9CA3AF
    TEXT = (55, 65, 81)          # #374151

    parsed = _parse_content(content)
    safe_name = _sanitize_filename(filename)
    legal = _is_legal(title, content)

    if preview:
        return _content_summary(parsed, "pdf")

    output_dir = os.path.realpath(CONFIG["output_dir"])
    os.makedirs(output_dir, exist_ok=True)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.set_margins(25, 20, 25)
    pdf.add_page()

    # Unicode font
    font_candidates = [
        ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]
    font_regular, font_bold = None, None
    for reg, bold in font_candidates:
        if os.path.exists(reg):
            font_regular, font_bold = reg, bold
            break
    if font_regular:
        pdf.add_font("DocFont", "", font_regular)
        if font_bold and os.path.exists(font_bold):
            pdf.add_font("DocFont", "B", font_bold)
        else:
            pdf.add_font("DocFont", "B", font_regular)
    else:
        pdf.add_font("DocFont", "", "helvetica")
        pdf.add_font("DocFont", "B", "helvetica")

    if not legal:
        # Minimal branding: thin orange top line
        pdf.set_draw_color(*ORANGE)
        pdf.set_line_width(1.2)
        pdf.line(25, 10, 185, 10)

        # Company name — top right, small, orange
        pdf.set_xy(25, 13)
        pdf.set_font("DocFont", "", 7.5)
        pdf.set_text_color(*GRAY)
        pdf.cell(160, 5, "the firm", align="R")
        pdf.ln(12)

        # Title — clean, large but not heavy
        pdf.set_font("DocFont", "B", 15)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

        # Date — below title, minimal
        pdf.set_font("DocFont", "", 8.5)
        pdf.set_text_color(*GRAY)
        pdf.cell(0, 5, datetime.now().strftime('%d.%m.%Y'), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(8)
    else:
        pdf.set_font("DocFont", "B", 14)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(1)
        pdf.set_font("DocFont", "", 8.5)
        pdf.set_text_color(*GRAY)
        pdf.cell(0, 5, datetime.now().strftime('%d.%m.%Y'), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(6)

    # Body — clean typography, minimal decoration
    lm = pdf.l_margin  # left margin reference
    for line_type, text in parsed:
        if line_type == "heading":
            pdf.ln(6)
            pdf.set_x(lm)
            pdf.set_font("DocFont", "B", 11)
            pdf.set_text_color(*DARK if not legal else (0, 0, 0))
            pdf.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
        elif line_type == "bullet":
            pdf.set_x(lm)
            pdf.set_font("DocFont", "", 10)
            pdf.set_text_color(*ORANGE if not legal else (100, 100, 100))
            pdf.cell(8, 6, "  \u2022")
            pdf.set_text_color(*TEXT if not legal else (0, 0, 0))
            pdf.cell(0, 6, f" {text}", new_x="LMARGIN", new_y="NEXT")
        elif line_type == "field":
            # T15: Key: Value — key bold, value regular
            key, val = text.split(":", 1)
            pdf.set_x(lm)
            pdf.set_font("DocFont", "B", 10)
            pdf.set_text_color(*DARK if not legal else (0, 0, 0))
            pdf.cell(pdf.get_string_width(key + ": ") + 2, 6, key + ": ")
            pdf.set_font("DocFont", "", 10)
            pdf.set_text_color(*TEXT if not legal else (0, 0, 0))
            pdf.cell(0, 6, val, new_x="LMARGIN", new_y="NEXT")
        elif line_type == "text":
            pdf.set_x(lm)
            pdf.set_font("DocFont", "", 10)
            pdf.set_text_color(*TEXT if not legal else (0, 0, 0))
            pdf.multi_cell(0, 6, text)
        elif line_type == "blank":
            pdf.ln(4)

    # Footer — bottom of page, nearly invisible
    pdf.ln(12)
    pdf.set_font("DocFont", "", 7)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 4, "Accounting Firm" if not legal else "Accounting Firm", new_x="LMARGIN", new_y="NEXT", align="C")

    filepath = f"{output_dir}/{safe_name}.pdf"
    pdf.output(filepath)
    summary = _content_summary(parsed, "pdf")
    return f"PDF created: {safe_name}.pdf ({summary}) — /files/{safe_name}.pdf"


def _create_docx(title, content, filename, preview=False):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn

    title = _fix_encoding(title)
    content = _fix_encoding(content)

    ORANGE = RGBColor(0xF9, 0x73, 0x16)
    DARK_NAVY = RGBColor(0x1F, 0x29, 0x37)
    GRAY = RGBColor(0x6B, 0x72, 0x80)

    parsed = _parse_content(content)
    safe_name = _sanitize_filename(filename)
    legal = _is_legal(title, content)

    if preview:
        return _content_summary(parsed, "docx")

    output_dir = os.path.realpath(CONFIG["output_dir"])
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    if not legal:
        # Header — company name top right, orange
        header = doc.sections[0].header
        hp = header.paragraphs[0]
        hp.alignment = 2  # RIGHT
        hr = hp.add_run("Accounting Firm")
        hr.font.size = Pt(9)
        hr.font.color.rgb = ORANGE
        hr.font.bold = True

        # Footer
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.alignment = 1  # CENTER
        fr = fp.add_run("Accounting Firm")
        fr.font.size = Pt(8)
        fr.font.color.rgb = GRAY

    # Title
    heading = doc.add_heading(level=0)
    run = heading.add_run(title)
    run.font.color.rgb = DARK_NAVY if not legal else RGBColor(0, 0, 0)

    if not legal:
        # Orange accent line (border bottom)
        pPr = heading._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '12',
            qn('w:space'): '1', qn('w:color'): 'F97316'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Date
    date_p = doc.add_paragraph()
    date_r = date_p.add_run(f"Date: {datetime.now().strftime('%d.%m.%Y')}")
    date_r.font.size = Pt(9)
    date_r.font.color.rgb = GRAY

    # Body
    for line_type, text in parsed:
        if line_type == "heading":
            h = doc.add_heading(level=2)
            hr = h.add_run(text)
            hr.font.color.rgb = DARK_NAVY if not legal else RGBColor(0, 0, 0)
        elif line_type == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif line_type == "field":
            # T15: Key: Value — key bold, value regular
            key, val = text.split(":", 1)
            p = doc.add_paragraph()
            rk = p.add_run(key + ": ")
            rk.bold = True
            rk.font.color.rgb = DARK_NAVY if not legal else RGBColor(0, 0, 0)
            rv = p.add_run(val)
            rv.font.color.rgb = DARK_NAVY if not legal else RGBColor(0, 0, 0)
        elif line_type == "text":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.color.rgb = DARK_NAVY if not legal else RGBColor(0, 0, 0)

    if legal:
        sig = doc.add_paragraph("\nAccounting Firm")
        sig.runs[0].font.size = Pt(9)

    filepath = f"{output_dir}/{safe_name}.docx"
    doc.save(filepath)
    summary = _content_summary(parsed, "docx")
    return f"DOCX created: {safe_name}.docx ({summary}) — /files/{safe_name}.docx"


def _create_xlsx(title, content, filename):
    import csv as csv_mod
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

    # Brand palette
    ORANGE_HEX = "F97316"
    DARK_NAVY_HEX = "1F2937"
    LIGHT_GRAY_HEX = "F3F4F6"
    GRAY_HEX = "E5E7EB"

    safe_name = _sanitize_filename(filename)
    output_dir = os.path.realpath(CONFIG["output_dir"])
    os.makedirs(output_dir, exist_ok=True)

    try:
        reader = csv_mod.reader(io.StringIO(content))
        rows = list(reader)
        if len(rows) < 1:
            return "Error: CSV data is empty."
        headers = rows[0]
        data_rows = rows[1:]
    except Exception as e:
        return f"Error: CSV parsing failed — {str(e)}"

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]

    thin_border = Border(
        left=Side(style='thin', color=GRAY_HEX),
        right=Side(style='thin', color=GRAY_HEX),
        top=Side(style='thin', color=GRAY_HEX),
        bottom=Side(style='thin', color=GRAY_HEX),
    )

    # Header row — orange
    header_fill = PatternFill(start_color=ORANGE_HEX, end_color=ORANGE_HEX, fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h.strip())
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # Data rows — alternating row colors
    alt_fill = PatternFill(start_color=LIGHT_GRAY_HEX, end_color=LIGHT_GRAY_HEX, fill_type="solid")
    data_font = Font(color=DARK_NAVY_HEX, size=10)
    for r, row_data in enumerate(data_rows, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val.strip())
            cell.font = data_font
            cell.border = thin_border
            if r % 2 == 0:
                cell.fill = alt_fill

    # Auto-width
    for col_cells in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col_cells)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 50)

    # Freeze header row
    ws.freeze_panes = "A2"

    filepath = f"{output_dir}/{safe_name}.xlsx"
    wb.save(filepath)
    return f"XLSX created: {safe_name}.xlsx ({len(data_rows)} rows, {len(headers)} columns) — /files/{safe_name}.xlsx"


def exa_search(query: str) -> str:
    """Performs a web search using the Exa AI search API.

    Searches for current information on accounting, tax, and legal topics.
    Requires the EXA_API_KEY environment variable.
    """
    import urllib.request
    import urllib.error
    import ssl

    api_key = os.getenv("EXA_API_KEY", "")
    if not api_key:
        return "Error: EXA_API_KEY environment variable is not set."

    url = "https://api.exa.ai/search"
    payload = json.dumps({
        "query": query,
        "type": "auto",
        "numResults": 5,
        "contents": {
            "text": {"maxCharacters": 500}
        }
    }).encode("utf-8")

    ctx = ssl.create_default_context()
    try:
        import certifi
        ctx.load_verify_locations(certifi.where())
    except ImportError:
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE

    try:
        req = urllib.request.Request(url, data=payload, headers={
            "x-api-key": api_key,
            "Content-Type": "application/json",
            "Accept": "application/json",
        })
        with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        return f"Exa API error: HTTP {e.code}"
    except Exception as e:
        return f"Exa API connection error: {str(e)[:100]}"

    results = data.get("results", [])
    if not results:
        return f"No results found for '{query}'."

    output = []
    for r in results[:5]:
        title = r.get("title", "?")
        url_r = r.get("url", "")
        text = r.get("text", "")[:300]
        output.append(f"**{title}**\n{url_r}\n{text}")

    return "\n---\n".join(output)


def krs_lookup(krs_number: str) -> str:
    """Queries company information from the Polish KRS (Krajowy Rejestr Sadowy) API.

    Free, public API. No authentication required. GDPR compliant (personal data anonymized).
    """
    import urllib.request
    import urllib.error
    import ssl

    # Clean KRS number — digits only
    krs_clean = "".join(c for c in str(krs_number) if c.isdigit())
    if not krs_clean:
        return f"Error: '{krs_number}' is not a valid KRS number."
    krs_clean = krs_clean.zfill(10)  # Pad to 10 digits

    url = f"https://api-krs.ms.gov.pl/api/krs/OdpisPelny/{krs_clean}?rejestr=P&format=json"

    # SSL context — for certificate issues on Windows
    ctx = ssl.create_default_context()
    try:
        import certifi
        ctx.load_verify_locations(certifi.where())
    except ImportError:
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE

    try:
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return f"KRS {krs_clean} not found. Please check the number."
        return f"KRS API error: HTTP {e.code}"
    except Exception as e:
        return f"KRS API connection error: {str(e)[:100]}"

    # Extract important fields from JSON structure — KRS API uses array-based format
    try:
        odpis = data.get("odpis", {})
        dane = odpis.get("dane", {})
        dzial1 = dane.get("dzial1", {})
        podmiot = dzial1.get("danePodmiotu", {})

        # nazwa, formaPrawna, identyfikatory are all lists
        nazwa_list = podmiot.get("nazwa", [])
        nazwa = nazwa_list[-1].get("nazwa", "?") if nazwa_list else "?"

        forma_list = podmiot.get("formaPrawna", [])
        forma = forma_list[-1].get("formaPrawna", "?") if forma_list else "?"

        id_list = podmiot.get("identyfikatory", [])
        ids = id_list[-1].get("identyfikatory", {}) if id_list else {}
        nip = ids.get("nip", "?")
        regon = ids.get("regon", "?")

        # Address
        siedziba_list = dzial1.get("siedzibaIAdres", {}).get("siedziba", [])
        if siedziba_list:
            s = siedziba_list[-1]
            adres_parts = [s.get("ulica", ""), s.get("nrDomu", ""), s.get("miejscowosc", ""), s.get("kodPocztowy", "")]
            adres = " ".join(p for p in adres_parts if p)
        else:
            adres = "?"

        # Alternative address — address field
        if adres == "?":
            adres_list = dzial1.get("siedzibaIAdres", {}).get("adres", [])
            if adres_list:
                a = adres_list[-1] if isinstance(adres_list, list) else adres_list
                if isinstance(a, dict):
                    adres_parts = [a.get("ulica", ""), a.get("nrDomu", ""), a.get("miejscowosc", ""), a.get("kodPocztowy", "")]
                    adres = " ".join(p for p in adres_parts if p)

        return (
            f"KRS: {krs_clean}\n"
            f"Company: {nazwa}\n"
            f"Forma prawna: {forma}\n"
            f"NIP: {nip}\n"
            f"REGON: {regon}\n"
            f"Address: {adres}"
        )
    except Exception:
        return json.dumps(data, ensure_ascii=False, indent=2)[:2000]


def execute_tool(name, args):
    tools = {
        "wiki_read": lambda a: wiki_read(a["article_path"]),
        "mastersheet_read": lambda a: mastersheet_read(a["query"]),
        "wiki_write": lambda a: wiki_write(a["filename"], a["content"]),
        "create_pdf": lambda a: _create_pdf(a["title"], a["content"], a["filename"], a.get("preview", False)),
        "create_docx": lambda a: _create_docx(a["title"], a["content"], a["filename"], a.get("preview", False)),
        "create_xlsx": lambda a: _create_xlsx(a["title"], a["content"], a["filename"]),
        "krs_lookup": lambda a: krs_lookup(a["krs_number"]),
        "exa_search": lambda a: exa_search(a["query"]),
    }
    fn = tools.get(name)
    if fn:
        try:
            return fn(args)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Error: {str(e)}"

    # MCP fallback — route to the server that owns this tool
    mcp_server = _mcp_tool_registry.get(name)
    if mcp_server and mcp_server in _mcp_processes:
        proc = _mcp_processes[mcp_server]
        if proc.poll() is None:
            return _call_mcp_tool(mcp_server, name, args)

    return f"Error: '{name}' unrecognized tool."


# =============================================================================
# AGENT LOOP
# =============================================================================


def classify_tier(query_type, tool_calls, routing_decision):
    """HiTL tier classification. Returns (tier, reason) tuple.

    Tier 1 (Autonomous): Information queries, wiki reads, status checks — no approval needed.
    Tier 2 (Approval): Message drafts, document preparation, wiki writes — approval required.
    Tier 3 (Human): Complaints, contracts, legal, tax advice — direct to human.
    """
    # Tier 3: Topics outside the agent's authority — requests the organization rejects
    if query_type == "constraint":
        return 3, "out_of_scope"

    # Tier 2: Tool usage involving writing/generation
    write_tools = {"wiki_write", "create_pdf", "create_docx", "create_xlsx"}
    used_write = [tc["name"] for tc in tool_calls if tc["name"] in write_tools]
    if used_write:
        return 2, f"write_operation:{','.join(used_write)}"

    # Tier 2: If there is a routing decision (redirect to human)
    if routing_decision:
        return 1, f"routing:{routing_decision}"

    # Tier 1: Information queries, routing, read operations
    return 1, "info_query"


def classify_query(query, tool_calls, response_text):
    """Analyze the query and response to determine query_type and routing_decision."""
    q = query.lower()
    resp = (response_text or "").lower()

    # Routing detection — who was the response routed to?
    routing_map = {
        "gosia": "Gosia (Payroll)",
        "kasia": "Kasia (Accounting)",
        "liudmila": "Liudmila (Accounting)",
        "aleksandra": "Aleksandra (Documents)",
        "jakub": "Jakub (Legal)",
        "kaan": "Kaan Bey (Management)",
    }
    routing_decision = None
    for name, label in routing_map.items():
        if name in resp:
            routing_decision = label
            break

    # Query type classification
    has_document = any(tc["name"] in ("create_pdf", "create_docx", "create_xlsx") for tc in tool_calls)
    has_mastersheet = any(tc["name"] == "mastersheet_read" for tc in tool_calls)
    if has_document:
        query_type = "document"
    elif has_mastersheet:
        query_type = "customer"
    elif routing_decision:
        query_type = "routing"
    elif any(w in q for w in ["tarih", "süre", "deadline", "ne zaman", "kaçı", "vade"]):
        query_type = "deadline"
    elif any(w in q for w in ["tavsiye", "hesapla", "ücret", "fiyat", "rakip", "pesel", "askıya"]):
        query_type = "constraint"
    else:
        query_type = "procedure"

    return query_type, routing_decision


def grounding_check(response_text, tool_results):
    """Compare numerical/date claims in the response against wiki content.

    Does not make additional LLM calls — deterministic string matching.
    Not full verification, catches obvious inconsistencies.
    """
    import re
    if not tool_results or not response_text:
        return "no_sources"

    combined_sources = " ".join(tool_results).lower()
    resp_lower = response_text.lower()

    # Extract numbers from response (exclude small numbers)
    resp_numbers = set(re.findall(r'\b\d{2,}\b', resp_lower))
    source_numbers = set(re.findall(r'\b\d{2,}\b', combined_sources))

    # Numbers in response but not in sources -> potential hallucination
    ungrounded = resp_numbers - source_numbers
    if ungrounded:
        return "flagged"
    return "passed"


def mask_pii(text: str) -> str:
    """S1: PII masking — mask NIP, PESEL, REGON, bank account numbers.

    NIP (10 digits) -> first 3 + *** + last 3
    PESEL (11 digits) -> first 2 + ******* + last 2
    REGON (9/14 digits) -> first 3 + *** + last 3
    Bank account (26 digit PL IBAN) -> PL** **** **** **** **** **** ****
    """
    import re
    if not text or not isinstance(text, str):
        return text or ""

    # NIP: 10-digit number (may contain spaces or dashes: 774-000-14-54 or 7740001454)
    def _mask_nip(m):
        digits = re.sub(r'[\s\-]', '', m.group(0))
        if len(digits) == 10:
            return digits[:3] + "***" + digits[7:]
        return m.group(0)
    text = re.sub(r'\b\d{3}[\-\s]?\d{3}[\-\s]?\d{2}[\-\s]?\d{2}\b', _mask_nip, text)

    # Plain 10-digit NIP (without dashes/spaces)
    def _mask_plain_nip(m):
        d = m.group(0)
        return d[:3] + "***" + d[7:]
    text = re.sub(r'\b\d{10}\b', _mask_plain_nip, text)

    # PESEL: 11 digits
    def _mask_pesel(m):
        d = m.group(0)
        return d[:2] + "*******" + d[9:]
    text = re.sub(r'\b\d{11}\b', _mask_pesel, text)

    # PL IBAN: PL + 26 digits or 26-digit number
    text = re.sub(r'\bPL\d{26}\b', 'PL**************************', text)
    text = re.sub(r'\b\d{26}\b', '**************************', text)

    return text


def _mask_trace_fields(trace: dict) -> dict:
    """Mask PII-containing fields in the trace dict."""
    masked = dict(trace)
    # Mask query text
    if "query" in masked:
        masked["query"] = mask_pii(str(masked["query"]))
    # Mask error detail
    if "error_detail" in masked:
        masked["error_detail"] = mask_pii(str(masked["error_detail"]))
    # Mask PII in tool args
    if "tool_calls" in masked and isinstance(masked["tool_calls"], list):
        masked_calls = []
        for tc in masked["tool_calls"]:
            mtc = dict(tc)
            if "args" in mtc and isinstance(mtc["args"], dict):
                mtc["args"] = {k: mask_pii(str(v)) if isinstance(v, str) else v
                               for k, v in mtc["args"].items()}
            masked_calls.append(mtc)
        masked["tool_calls"] = masked_calls
    return masked


def save_trace(trace):
    os.makedirs(CONFIG["trace_dir"], exist_ok=True)
    filepath = f"{CONFIG['trace_dir']}/{datetime.now().strftime('%Y-%m-%d')}.jsonl"
    # S1: Mask PII before writing to trace
    masked_trace = _mask_trace_fields(trace)
    with open(filepath, "a", encoding="utf-8") as f:
        f.write(json.dumps(masked_trace, ensure_ascii=False) + "\n")
    # T5: Alerting — check conditions after each trace
    _check_alerts(masked_trace)


# T5: Alert mechanism — pure logic, no additional LLM calls
_alert_buffer = []  # Keeps last N traces for pattern detection


def _check_alerts(trace):
    """Check alert conditions after each trace."""
    _alert_buffer.append(trace)
    # Keep last 20 traces (memory limit)
    if len(_alert_buffer) > 20:
        _alert_buffer.pop(0)

    alerts = []
    # Condition 1: 3+ errors in last 5 traces
    recent = _alert_buffer[-5:]
    error_count = sum(1 for t in recent if t.get("query_type") == "error" or t.get("error_category"))
    if error_count >= 3:
        alerts.append({"type": "high_error_rate", "detail": f"{error_count} errors in last 5 queries", "severity": "critical"})

    # Condition 2: Average latency > 15s in last 5 traces
    durations = [t.get("duration_ms", 0) for t in recent if t.get("duration_ms")]
    if durations and sum(durations) / len(durations) > 15000:
        avg = sum(durations) / len(durations) / 1000
        alerts.append({"type": "high_latency", "detail": f"Average latency: {avg:.1f}s", "severity": "warning"})

    # Condition 3: Same tool failing 3+ times consecutively
    if len(_alert_buffer) >= 3:
        last3_errors = []
        for t in _alert_buffer[-3:]:
            if t.get("error_category"):
                for tc in t.get("tool_calls", []):
                    last3_errors.append(tc.get("name"))
        if last3_errors and len(set(last3_errors)) == 1 and len(last3_errors) >= 3:
            alerts.append({"type": "tool_failure", "detail": f"Tool '{last3_errors[0]}' consecutive failures", "severity": "critical"})

    # Condition 4: High grounding check failure rate
    grounding_fails = sum(1 for t in _alert_buffer if t.get("grounding_result") == "flagged")
    if len(_alert_buffer) >= 5 and grounding_fails / len(_alert_buffer) > 0.2:
        alerts.append({"type": "grounding_issue", "detail": f"Grounding fail: {grounding_fails}/{len(_alert_buffer)}", "severity": "warning"})

    # Save alerts
    if alerts:
        alerts_path = os.path.join(CONFIG["trace_dir"], "alerts.jsonl")
        for alert in alerts:
            alert["timestamp"] = datetime.now().isoformat()
            with open(alerts_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(alert, ensure_ascii=False) + "\n")


def model_call(messages, tools=None):
    # Anthropic API — automatic detection via key prefix
    if CONFIG["api_key"].startswith("sk-ant-"):
        return _anthropic_call(messages, tools)

    # OpenAI-compatible API (OpenRouter, LM Studio, etc.)
    client = get_client()
    kwargs = {
        "model": CONFIG["model"],
        "messages": messages,
        "temperature": CONFIG["temperature"],
        "max_tokens": CONFIG["max_tokens"],
        "timeout": 300,
    }
    if tools:
        kwargs["tools"] = tools
    for attempt in range(3):
        try:
            msg = client.chat.completions.create(**kwargs).choices[0].message
            # Thinking mode fix: some models write to reasoning_content, leaving content empty
            if not msg.content and not msg.tool_calls:
                reasoning = getattr(msg, 'reasoning_content', None)
                if reasoning:
                    msg.content = reasoning
            return msg
        except Exception as e:
            if "429" in str(e) and attempt < 2:
                time.sleep(15 * (attempt + 1))
                continue
            raise


def _anthropic_call(messages, tools=None):
    """Anthropic Messages API — converts OpenAI-format messages and calls the API."""
    from anthropic import Anthropic
    from types import SimpleNamespace

    client = Anthropic(api_key=CONFIG["api_key"])

    # Separate system message, convert others to Anthropic format
    system_text = ""
    api_messages = []

    for m in messages:
        role = m["role"]

        if role == "system":
            system_text += ("\n\n" if system_text else "") + (m.get("content") or "")

        elif role == "assistant":
            if m.get("tool_calls"):
                content = []
                if m.get("content"):
                    content.append({"type": "text", "text": m["content"]})
                for tc in m["tool_calls"]:
                    tc_id = tc.id if hasattr(tc, "id") else tc.get("id", "")
                    tc_name = tc.function.name if hasattr(tc, "function") else tc.get("function", {}).get("name", "")
                    tc_args = tc.function.arguments if hasattr(tc, "function") else tc.get("function", {}).get("arguments", "{}")
                    content.append({
                        "type": "tool_use",
                        "id": tc_id,
                        "name": tc_name,
                        "input": json.loads(tc_args) if isinstance(tc_args, str) else tc_args,
                    })
                api_messages.append({"role": "assistant", "content": content})
            else:
                api_messages.append({"role": "assistant", "content": m.get("content") or ""})

        elif role == "tool":
            tool_result = {
                "type": "tool_result",
                "tool_use_id": m["tool_call_id"],
                "content": m["content"],
            }
            # Merge consecutive tool_results into a single user message
            if api_messages and api_messages[-1]["role"] == "user" and isinstance(api_messages[-1]["content"], list):
                api_messages[-1]["content"].append(tool_result)
            else:
                api_messages.append({"role": "user", "content": [tool_result]})

        else:  # user
            api_messages.append({"role": "user", "content": m.get("content") or ""})

    # Tools: convert from OpenAI format to Anthropic format
    api_tools = None
    if tools:
        api_tools = [
            {
                "name": t["function"]["name"],
                "description": t["function"]["description"],
                "input_schema": t["function"]["parameters"],
            }
            for t in tools
        ]

    kwargs = {
        "model": CONFIG["model"],
        "messages": api_messages,
        "system": system_text,
        "temperature": CONFIG["temperature"],
        "max_tokens": CONFIG["max_tokens"],
    }
    if api_tools:
        kwargs["tools"] = api_tools

    for attempt in range(3):
        try:
            response = client.messages.create(**kwargs)
            break
        except Exception as e:
            if "429" in str(e) and attempt < 2:
                time.sleep(15 * (attempt + 1))
                continue
            raise

    # Convert response to OpenAI-compatible SimpleNamespace object
    text_parts = []
    tool_calls = []

    for block in response.content:
        if block.type == "text":
            text_parts.append(block.text)
        elif block.type == "tool_use":
            tool_calls.append(SimpleNamespace(
                id=block.id,
                function=SimpleNamespace(
                    name=block.name,
                    arguments=json.dumps(block.input, ensure_ascii=False),
                ),
            ))

    return SimpleNamespace(
        content=" ".join(text_parts) if text_parts else None,
        tool_calls=tool_calls if tool_calls else None,
    )


def estimate_tokens(text):
    """Rough token estimate — chars/3 approach for Turkish/Polish text."""
    if not text:
        return 0
    return len(text) // 3


def build_messages(session):
    index_content = read_file(CONFIG["index_file"])
    max_index = CONFIG.get("max_index_chars", 4000)
    if len(index_content) > max_index:
        index_content = index_content[:max_index] + "\n\n... [INDEX truncated — use wiki_read('INDEX.md') for more articles]"
    system = SYSTEM_PROMPT + f"\n\n[WIKI INDEX — Decide which article to read by looking at this list]\n{index_content}"
    messages = [{"role": "system", "content": system}]
    messages.extend(session["messages"][-CONFIG["max_history_messages"]:])

    # Token estimation — context overflow safety net
    total_tokens = sum(estimate_tokens(m.get("content", "")) for m in messages)
    context_limit = 28000  # Qwen 3.6 32K - leave room for max_tokens(4096)
    if total_tokens > context_limit:
        # Trim old messages from history
        while total_tokens > context_limit and len(messages) > 2:
            removed = messages.pop(1)  # Remove oldest message after system
            total_tokens -= estimate_tokens(removed.get("content", ""))

    return messages


def _friendly_error(e: Exception) -> str:
    """T2: Return API errors as user-friendly messages."""
    err = str(e).lower()
    if "429" in err or "rate" in err:
        return "API rate limit exceeded. Please try again in 1-2 minutes."
    if "timeout" in err or "timed out" in err:
        return "Response timed out. Try shortening your question."
    if "connection" in err or "connect" in err:
        return "Could not connect to server. Please check your internet connection."
    if "401" in err or "auth" in err:
        return "API authentication failed. Please contact the administrator."
    if "500" in err or "server" in err:
        return "An error occurred on the remote server. Please try again shortly."
    return "An error occurred. Please try again."


def agent_loop(user_message: str, session: dict) -> str:
    """think → call → observe → respond → trace"""
    start = time.time()

    # S4: Input sanitization — max length at agent_loop entry
    max_len = CONFIG.get("max_query_length", 2000)
    if len(user_message) > max_len:
        user_message = user_message[:max_len]

    session["messages"].append({"role": "user", "content": user_message})
    messages = build_messages(session)

    tool_calls_log = []
    wiki_articles_used = []
    tool_call_count = 0

    # T2: Wrap entire agent loop in try/except — message instead of crash on error
    try:
        return _agent_loop_inner(user_message, session, messages, start)
    except Exception as e:
        error_msg = _friendly_error(e)
        session["messages"].append({"role": "assistant", "content": error_msg})
        save_trace({
            "timestamp": datetime.now().isoformat(),
            "query": user_message[:200],
            "query_type": "error",
            "tier": 0,
            "tool_calls": [],
            "error_category": type(e).__name__,
            "error_detail": str(e)[:500],
            "duration_ms": int((time.time() - start) * 1000),
        })
        return error_msg


def _agent_loop_inner(user_message, session, messages, start):
    """Agent loop inner logic — wrapped by agent_loop."""
    tool_calls_log = []
    wiki_articles_used = []
    tool_call_count = 0
    turn_count = 0  # Number of model calls (turns), not tool calls

    while turn_count < CONFIG["max_tool_calls_per_turn"]:
        response = model_call(messages, tools=TOOLS)

        if response.tool_calls:
            turn_count += 1  # One turn = one model call (regardless of tool count)
            # Group all tool calls in a single assistant message (OpenAI spec)
            messages.append({
                "role": "assistant",
                "content": response.content,
                "tool_calls": list(response.tool_calls),
            })
            for tc in response.tool_calls:
                tool_call_count += 1
                args = json.loads(tc.function.arguments)
                result = execute_tool(tc.function.name, args)
                tool_calls_log.append({"name": tc.function.name, "args": args})
                if tc.function.name == "wiki_read":
                    wiki_articles_used.append(args.get("article_path", ""))
                messages.append({"role": "tool", "tool_call_id": tc.id, "content": result})
            continue

        # None guard — model may sometimes return content=None
        final = response.content or "Could not generate a response. Please try rephrasing your question."
        session["messages"].append({"role": "assistant", "content": final})

        # Multi-turn tool context — previous tool results accessible in next turn
        if tool_calls_log:
            tool_summary_parts = []
            for tc in tool_calls_log:
                if tc["name"] == "mastersheet_read":
                    tool_summary_parts.append(f"mastersheet({tc['args'].get('query', '')})")
                elif tc["name"] == "wiki_read":
                    tool_summary_parts.append(f"wiki({tc['args'].get('article_path', '')})")
            if tool_summary_parts:
                session["messages"].append({
                    "role": "system",
                    "content": f"[Previous turn tools: {', '.join(tool_summary_parts)}]"
                })

        # Session memory trimming — prevent unlimited growth
        max_keep = CONFIG["max_history_messages"] * 2
        if len(session["messages"]) > max_keep:
            session["messages"] = session["messages"][-max_keep:]

        # Automatic classification
        query_type, routing_decision = classify_query(user_message, tool_calls_log, final)

        # Self-check: source tag + grounding verification
        has_source = "[SOURCE:" in final or "[KAYNAK:" in final or "I cannot provide definitive information" in final or "Bu konuda kesin bilgi veremiyorum" in final
        tool_result_texts = [m["content"] for m in messages if m.get("role") == "tool"]
        grounding = grounding_check(final, tool_result_texts)

        # Routing responses don't need wiki — routing matrix in system prompt suffices
        is_routing = query_type == "routing" or routing_decision is not None

        if not has_source and not is_routing:
            self_check_result = "missing_source"
        elif grounding == "flagged":
            self_check_result = "ungrounded_claim"
        else:
            self_check_result = "passed"

        # T9: classify_tier now returns (tier, reason) tuple
        tier, tier_reason = classify_tier(query_type, tool_calls_log, routing_decision)

        # T9: Save tool usage summary as reasoning
        tool_reasoning = ", ".join(
            f"{tc['name']}({list(tc['args'].values())[0][:30] if tc['args'] else ''})"
            for tc in tool_calls_log
        ) if tool_calls_log else "no_tools_used"

        trace_data = {
            "timestamp": datetime.now().isoformat(),
            "query": user_message,
            "query_type": query_type,
            "tier": tier,
            "tier_reason": tier_reason,
            "tool_reasoning": tool_reasoning,
            "wiki_articles_used": wiki_articles_used,
            "tool_calls": tool_calls_log,
            "routing_decision": routing_decision,
            "self_check_result": self_check_result,
            "grounding_result": grounding,
            "error_category": None,
            "duration_ms": int((time.time() - start) * 1000),
        }
        save_trace(trace_data)

        # T10: Update client state (for clients queried from mastersheet)
        try:
            _update_client_state_from_trace(trace_data)
        except Exception:
            pass  # State update error should not break the agent

        return final

    fallback = "Too many steps required. Please ask a more specific question."
    session["messages"].append({"role": "assistant", "content": fallback})
    return fallback


# =============================================================================
# CLI
# =============================================================================

# =============================================================================
# HTTP API
# =============================================================================

_sessions = {}


# =============================================================================
# T3: SESSION PERSISTENCE — write/read to disk, no loss after restart
# =============================================================================

def _session_path(session_id: str) -> str:
    """Return the disk path for a session file."""
    import re
    safe_id = re.sub(r'[^\w\-]', '_', session_id)[:64]
    return os.path.join(CONFIG.get("session_dir", "traces/sessions"), f"{safe_id}.json")


def _save_session(session_id: str, session: dict):
    """Save session to disk. tool_calls objects in messages are made JSON-safe."""
    session_dir = CONFIG.get("session_dir", "traces/sessions")
    os.makedirs(session_dir, exist_ok=True)

    # Make messages JSON-serializable (tool_calls objects must be converted to dicts)
    safe_messages = []
    for msg in session.get("messages", []):
        m = dict(msg)
        if "tool_calls" in m and m["tool_calls"]:
            m["tool_calls"] = [
                {
                    "id": getattr(tc, "id", tc.get("id", "")) if hasattr(tc, "id") else tc.get("id", ""),
                    "function": {
                        "name": (tc.function.name if hasattr(tc, "function") and hasattr(tc.function, "name")
                                 else tc.get("function", {}).get("name", "")),
                        "arguments": (tc.function.arguments if hasattr(tc, "function") and hasattr(tc.function, "arguments")
                                      else tc.get("function", {}).get("arguments", "{}")),
                    }
                }
                for tc in m["tool_calls"]
            ]
        safe_messages.append(m)

    data = {
        "session_id": session_id,
        "messages": safe_messages[-CONFIG.get("max_history_messages", 10) * 2:],
        "last_active": datetime.now().isoformat(),
    }
    filepath = _session_path(session_id)
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)


def _load_sessions():
    """Load session files from disk. Archive expired ones."""
    session_dir = CONFIG.get("session_dir", "traces/sessions")
    if not os.path.exists(session_dir):
        return

    expire_hours = CONFIG.get("session_expire_hours", 24)
    from datetime import timedelta
    cutoff = (datetime.now() - timedelta(hours=expire_hours)).isoformat()

    loaded = 0
    archived = 0
    max_sessions = CONFIG.get("max_active_sessions", 50)

    for fname in sorted(os.listdir(session_dir)):
        if not fname.endswith(".json"):
            continue
        filepath = os.path.join(session_dir, fname)
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            sid = data.get("session_id", fname[:-5])
            last_active = data.get("last_active", "")

            # Expiration check
            if last_active and last_active < cutoff:
                # Archive it
                archive_dir = os.path.join(session_dir, "archived")
                os.makedirs(archive_dir, exist_ok=True)
                os.rename(filepath, os.path.join(archive_dir, fname))
                archived += 1
                continue

            # Max session check
            if loaded >= max_sessions:
                continue

            _sessions[sid] = {"messages": data.get("messages", [])}
            loaded += 1
        except (json.JSONDecodeError, KeyError):
            continue

    if loaded > 0 or archived > 0:
        print(f"[T3] Session persistence: {loaded} loaded, {archived} archived")


# =============================================================================
# T10: CLIENT STATE — persistent client state (LeCun: World Model)
# =============================================================================

_CLIENT_STATE_PATH = os.path.join(CONFIG.get("trace_dir", "traces"), "client_state.json")
_client_state = {}


def _load_client_state():
    """Load client state JSON from disk."""
    global _client_state
    if os.path.exists(_CLIENT_STATE_PATH):
        try:
            with open(_CLIENT_STATE_PATH, "r", encoding="utf-8") as f:
                _client_state = json.load(f)
        except (json.JSONDecodeError, IOError):
            _client_state = {}


def _save_client_state():
    """Write client state to disk."""
    os.makedirs(os.path.dirname(_CLIENT_STATE_PATH), exist_ok=True)
    with open(_CLIENT_STATE_PATH, "w", encoding="utf-8") as f:
        json.dump(_client_state, f, ensure_ascii=False, indent=2)


def _init_client_state():
    """Create initial state for all clients from mastersheet (if not exists)."""
    import csv
    _load_client_state()
    try:
        with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                nip = safe_nip(row.get("NIP/PESEL", ""))
                name = row.get("Company Name", "").strip()
                if not nip or not name:
                    continue
                if nip not in _client_state:
                    _client_state[nip] = {
                        "company": name,
                        "typ": row.get("TYP", ""),
                        "responsible": row.get("RA", ""),
                        "last_contact": None,
                        "total_queries": 0,
                        "last_query_topic": None,
                        "risk_flags": [],
                    }
    except FileNotFoundError:
        pass
    _save_client_state()


def _update_client_state_from_trace(trace: dict):
    """Update relevant client state after agent loop."""
    # Try to extract client NIP/name from tool calls
    for tc in trace.get("tool_calls", []):
        if tc.get("name") == "mastersheet_read":
            query = tc.get("args", {}).get("query", "")
            # Try exact match by NIP
            nip_clean = query.strip()
            try:
                nip_clean = str(int(float(nip_clean)))
            except (ValueError, TypeError):
                pass
            # Look up NIP in state
            if nip_clean in _client_state:
                state = _client_state[nip_clean]
                state["last_contact"] = trace.get("timestamp")
                state["total_queries"] = state.get("total_queries", 0) + 1
                state["last_query_topic"] = trace.get("query_type")
                _save_client_state()
                return
            # Fuzzy match by company name
            q_lower = query.lower().strip()
            for nip, state in _client_state.items():
                if q_lower and q_lower in state.get("company", "").lower():
                    state["last_contact"] = trace.get("timestamp")
                    state["total_queries"] = state.get("total_queries", 0) + 1
                    state["last_query_topic"] = trace.get("query_type")
                    _save_client_state()
                    return


def create_api():
    from fastapi import FastAPI, Request
    from fastapi.responses import JSONResponse
    from pydantic import BaseModel

    # --- S3: Rate Limiting ---
    try:
        from slowapi import Limiter, _rate_limit_exceeded_handler
        from slowapi.util import get_remote_address
        from slowapi.errors import RateLimitExceeded

        limiter = Limiter(key_func=get_remote_address)
        _has_limiter = True
    except ImportError:
        limiter = None
        _has_limiter = False

    app = FastAPI(title="CS Agent", version="1.0")

    if _has_limiter:
        app.state.limiter = limiter

        @app.exception_handler(RateLimitExceeded)
        async def rate_limit_handler(request: Request, exc: RateLimitExceeded):
            return JSONResponse(
                status_code=429,
                content={"error": True, "message": "Too many requests. Please wait a moment."}
            )

    class ChatRequest(BaseModel):
        message: str
        session_id: str = "default"

    class FeedbackRequest(BaseModel):
        session_id: str = "default"
        rating: int  # 1 = thumbs up, -1 = thumbs down
        comment: str = ""
        message_index: int = -1  # T1: Which message this belongs to (-1 = last message)

    # --- S3: Rate limit decorator helper ---
    def _limit(rate_key):
        """Rate limit decorator — no-op if slowapi is not installed."""
        if _has_limiter:
            return limiter.limit(CONFIG.get(rate_key, "30/minute"))
        return lambda fn: fn

    @app.post("/chat")
    @_limit("rate_limit_chat")
    def chat(req: ChatRequest, request: Request):
        # --- S4: Input sanitization ---
        max_len = CONFIG.get("max_query_length", 2000)
        if len(req.message) > max_len:
            return {"error": True, "message": f"Message too long (max {max_len} characters). Please shorten and try again."}
        if not req.message.strip():
            return {"error": True, "message": "Empty message cannot be sent."}

        if req.session_id not in _sessions:
            _sessions[req.session_id] = {"messages": []}
        session = _sessions[req.session_id]

        # --- T2: Error handling — API errors returned as user-friendly messages ---
        try:
            response = agent_loop(req.message, session)
            # T3: Save session to disk after each message
            _save_session(req.session_id, session)
            return {"response": response, "session_id": req.session_id}
        except Exception as e:
            error_msg = _friendly_error(e)
            _save_session(req.session_id, session)
            # Log error to trace
            save_trace({
                "timestamp": datetime.now().isoformat(),
                "query": req.message[:200],
                "query_type": "error",
                "tier": 0,
                "tool_calls": [],
                "error_category": type(e).__name__,
                "error_detail": str(e)[:500],
                "duration_ms": 0,
            })
            return {"error": True, "message": error_msg, "session_id": req.session_id}

    @app.post("/feedback")
    @_limit("rate_limit_feedback")
    def feedback(req: FeedbackRequest, request: Request):
        # --- S4: Comment length limit ---
        max_comment = CONFIG.get("max_feedback_comment_length", 500)
        comment = req.comment[:max_comment] if req.comment else ""

        # T1: Link feedback to message — which message it belongs to and what was asked
        linked_query = ""
        if req.session_id in _sessions:
            msgs = _sessions[req.session_id].get("messages", [])
            # Match by message_index, -1 means last assistant message
            idx = req.message_index
            if idx == -1:
                # Find the user message before the last assistant message
                for i in range(len(msgs) - 1, -1, -1):
                    if msgs[i].get("role") == "user":
                        linked_query = msgs[i].get("content", "")[:200]
                        idx = i
                        break
            elif 0 <= idx < len(msgs):
                # Find the user message before the specified index
                for i in range(min(idx, len(msgs) - 1), -1, -1):
                    if msgs[i].get("role") == "user":
                        linked_query = msgs[i].get("content", "")[:200]
                        break

        feedback_entry = {
            "timestamp": datetime.now().isoformat(),
            "session_id": req.session_id,
            "message_index": req.message_index,
            "linked_query": mask_pii(linked_query),
            "rating": req.rating,
            "comment": comment,
        }
        os.makedirs(CONFIG["trace_dir"], exist_ok=True)
        filepath = f"{CONFIG['trace_dir']}/feedback.jsonl"
        with open(filepath, "a", encoding="utf-8") as f:
            f.write(json.dumps(feedback_entry, ensure_ascii=False) + "\n")
        return {"status": "ok"}

    @app.get("/health")
    def health():
        return {"status": "ok", "model": CONFIG["model"]}

    @app.get("/files/{filename}")
    def download_file(filename: str):
        import re as re_mod
        safe = re_mod.sub(r'[^\w.\-]', '_', filename)
        output_dir = os.path.realpath(CONFIG["output_dir"])
        filepath = os.path.realpath(f"{CONFIG['output_dir']}/{safe}")
        if not filepath.startswith(output_dir) or not os.path.exists(filepath):
            from fastapi.responses import JSONResponse
            return JSONResponse(status_code=404, content={"error": "File not found"})
        from fastapi.responses import FileResponse
        return FileResponse(filepath, filename=safe)

    @app.get("/")
    def chat_ui():
        from fastapi.responses import HTMLResponse
        html_path = os.path.join(os.path.dirname(__file__) or ".", "chat.html")
        if os.path.exists(html_path):
            with open(html_path, "r", encoding="utf-8") as f:
                return HTMLResponse(content=f.read())
        return HTMLResponse(content="<h1>chat.html not found</h1>", status_code=404)

    @app.get("/favicon.ico")
    def favicon():
        from fastapi.responses import FileResponse
        ico_path = os.path.join(os.path.dirname(__file__) or ".", "favicon.ico")
        if os.path.exists(ico_path):
            return FileResponse(ico_path, media_type="image/x-icon")
        from fastapi.responses import Response
        return Response(status_code=204)

    # --- Admin API ---

    @app.get("/admin/stats")
    @_limit("rate_limit_admin")
    def admin_stats(request: Request):
        traces = load_traces()
        if not traces:
            return {"error": "No trace data"}
        total = len(traces)
        tiers = {1: 0, 2: 0, 3: 0}
        types = {}
        for t in traces:
            tiers[t.get("tier", 1)] = tiers.get(t.get("tier", 1), 0) + 1
            qt = t.get("query_type", "?")
            types[qt] = types.get(qt, 0) + 1
        return {"total": total, "tiers": tiers, "query_types": types}

    @app.get("/admin/report")
    @_limit("rate_limit_admin")
    def admin_report(request: Request):
        traces = load_traces(days=7)
        if not traces:
            return {"error": "No traces in the last 7 days"}
        total = len(traces)
        tiers = {1: 0, 2: 0, 3: 0}
        types = {}
        checks = {}
        durations = []
        wiki_usage = {}
        routing_targets = {}
        for t in traces:
            tiers[t.get("tier", 1)] = tiers.get(t.get("tier", 1), 0) + 1
            qt = t.get("query_type", "?")
            types[qt] = types.get(qt, 0) + 1
            sc = t.get("self_check_result", "passed")
            checks[sc] = checks.get(sc, 0) + 1
            if t.get("duration_ms"):
                durations.append(t["duration_ms"])
            for art in t.get("wiki_articles_used", []):
                wiki_usage[art] = wiki_usage.get(art, 0) + 1
            rd = t.get("routing_decision")
            if rd:
                routing_targets[rd] = routing_targets.get(rd, 0) + 1
        avg_ms = sum(durations) / len(durations) if durations else 0
        return {
            "total": total, "avg_latency_ms": round(avg_ms),
            "tiers": tiers, "query_types": types, "self_checks": checks,
            "top_wiki": dict(sorted(wiki_usage.items(), key=lambda x: -x[1])[:5]),
            "routing_targets": routing_targets,
        }

    @app.get("/admin/deadlines")
    @_limit("rate_limit_admin")
    def admin_deadlines(request: Request, ahead: int = 5):
        from datetime import timedelta
        today = datetime.now()
        target = today + timedelta(days=ahead)
        monthly = {
            7: "ZUS — small firms", 10: "ZUS — JDG",
            15: "ZUS — medium/large", 20: "PIT/CIT advance", 25: "JPK_V7M / VAT",
        }
        annual = {
            (1, 31): "PIT-11 to employees", (2, 28): "PIT-11 to tax office",
            (4, 30): "PIT-36/37 annual", (6, 30): "CIT-8 corporate",
        }
        upcoming = []
        for day, desc in monthly.items():
            try:
                d = today.replace(day=day)
            except ValueError:
                continue
            if today.date() <= d.date() <= target.date():
                upcoming.append({"date": str(d.date()), "days_left": (d.date() - today.date()).days, "desc": desc})
        for (month, day), desc in annual.items():
            try:
                d = today.replace(month=month, day=day)
            except ValueError:
                continue
            if today.date() <= d.date() <= target.date():
                upcoming.append({"date": str(d.date()), "days_left": (d.date() - today.date()).days, "desc": desc})
        upcoming.sort(key=lambda x: x["date"])
        return {"today": str(today.date()), "ahead_days": ahead, "deadlines": upcoming}

    @app.get("/admin/silent")
    @_limit("rate_limit_admin")
    def admin_silent(request: Request, days: int = 30):
        import csv as csv_mod
        traces = load_traces()
        last_contact = {}
        for t in traces:
            for tc in t.get("tool_calls", []):
                if tc.get("name") == "mastersheet_read":
                    q = tc.get("args", {}).get("query", "").lower()
                    ts = t.get("timestamp", "")
                    if q and ts and (q not in last_contact or ts > last_contact[q]):
                        last_contact[q] = ts
        companies = []
        try:
            with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
                for row in csv_mod.DictReader(f):
                    name = row.get("Company Name", "").strip()
                    if name:
                        companies.append(name)
        except FileNotFoundError:
            return {"error": "Mastersheet not found"}
        from datetime import timedelta
        cutoff = (datetime.now() - timedelta(days=days)).isoformat()
        silent = []
        for company in companies:
            cl = company.lower()
            matched_ts = None
            for q, ts in last_contact.items():
                if q in cl or cl in q:
                    if matched_ts is None or ts > matched_ts:
                        matched_ts = ts
            if matched_ts is None or matched_ts < cutoff:
                silent.append({"company": company, "last_contact": matched_ts})
        return {"days": days, "total_companies": len(companies), "silent_count": len(silent), "silent": silent[:50]}

    @app.get("/admin/gaps")
    @_limit("rate_limit_admin")
    def admin_gaps(request: Request):
        traces = load_traces()
        from collections import Counter
        gaps = [t.get("query", "?") for t in traces
                if t.get("self_check_result") == "missing_source" and t.get("query_type") != "constraint"]
        no_wiki = [t.get("query", "?") for t in traces
                   if t.get("query_type") in ("procedure", "deadline") and not t.get("wiki_articles_used")]
        return {
            "missing_source": Counter(gaps).most_common(10),
            "no_wiki_used": Counter(no_wiki).most_common(10),
        }

    # --- T5: Alerts API ---
    @app.get("/admin/alerts")
    @_limit("rate_limit_admin")
    def admin_alerts(request: Request, hours: int = 24):
        alerts_path = os.path.join(CONFIG["trace_dir"], "alerts.jsonl")
        if not os.path.exists(alerts_path):
            return {"alerts": [], "total": 0}
        from datetime import timedelta
        cutoff = (datetime.now() - timedelta(hours=hours)).isoformat()
        alerts = []
        with open(alerts_path, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    a = json.loads(line)
                    if a.get("timestamp", "") >= cutoff:
                        alerts.append(a)
                except json.JSONDecodeError:
                    continue
        return {"alerts": alerts[-50:], "total": len(alerts)}

    # --- T4: Monitoring Dashboard ---
    @app.get("/admin/dashboard")
    @_limit("rate_limit_admin")
    def admin_dashboard(request: Request):
        from fastapi.responses import HTMLResponse
        traces = load_traces(days=1)
        total = len(traces) if traces else 0
        errors = sum(1 for t in traces if t.get("query_type") == "error" or t.get("error_category")) if traces else 0
        error_rate = (errors / total * 100) if total > 0 else 0
        durations = [t["duration_ms"] for t in traces if t.get("duration_ms")] if traces else []
        avg_ms = sum(durations) / len(durations) if durations else 0
        tiers = {1: 0, 2: 0, 3: 0}
        types = {}
        wiki_usage = {}
        for t in (traces or []):
            tiers[t.get("tier", 1)] = tiers.get(t.get("tier", 1), 0) + 1
            qt = t.get("query_type", "?")
            types[qt] = types.get(qt, 0) + 1
            for art in t.get("wiki_articles_used", []):
                wiki_usage[art] = wiki_usage.get(art, 0) + 1
        top_wiki = sorted(wiki_usage.items(), key=lambda x: -x[1])[:5]
        top_types = sorted(types.items(), key=lambda x: -x[1])[:6]
        # Alert status
        alert_class = "alert-red" if error_rate > 5 or avg_ms > 10000 else "alert-green"
        alert_text = "Warning: High error rate!" if error_rate > 5 else ("Warning: High latency!" if avg_ms > 10000 else "System normal")
        # Wiki rows
        wiki_rows = "".join(f"<tr><td>{a}</td><td>{c}</td></tr>" for a, c in top_wiki) if top_wiki else "<tr><td colspan=2>No data</td></tr>"
        type_rows = "".join(f"<tr><td>{t}</td><td>{c}</td></tr>" for t, c in top_types) if top_types else "<tr><td colspan=2>No data</td></tr>"
        html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><title>CS Agent Dashboard</title>
<meta http-equiv="refresh" content="30">
<style>
*{{margin:0;padding:0;box-sizing:border-box}}body{{font-family:system-ui;background:#f5f5f5;padding:24px;color:#333}}
h1{{color:#ff6600;margin-bottom:20px;font-size:22px}}
.grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:16px;margin-bottom:24px}}
.card{{background:#fff;border-radius:12px;padding:20px;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
.card h3{{font-size:13px;color:#888;text-transform:uppercase;margin-bottom:8px}}
.card .val{{font-size:28px;font-weight:700}}
.alert-green .val{{color:#22c55e}}.alert-red .val{{color:#ef4444}}
table{{width:100%;border-collapse:collapse;margin-top:8px}}th,td{{text-align:left;padding:8px 12px;border-bottom:1px solid #eee}}
th{{color:#888;font-size:12px;text-transform:uppercase}}
.section{{background:#fff;border-radius:12px;padding:20px;box-shadow:0 1px 3px rgba(0,0,0,.1);margin-bottom:16px}}
.section h2{{font-size:16px;margin-bottom:12px;color:#ff6600}}
.bar{{display:inline-block;height:20px;background:#ff6600;border-radius:4px;margin-right:8px}}
</style></head><body>
<h1>CS Agent — Dashboard</h1>
<div class="grid">
<div class="card"><h3>Queries (24h)</h3><div class="val">{total}</div></div>
<div class="card"><h3>Avg. Latency</h3><div class="val">{avg_ms/1000:.1f}s</div></div>
<div class="card"><h3>Error Rate</h3><div class="val" style="color:{'#ef4444' if error_rate>5 else '#22c55e'}">{error_rate:.1f}%</div></div>
<div class="card {alert_class}"><h3>Status</h3><div class="val">{alert_text}</div></div>
</div>
<div class="grid" style="grid-template-columns:1fr 1fr">
<div class="section"><h2>Tier Distribution</h2>
<div>Tier 1 (Autonomous): <strong>{tiers[1]}</strong> <div class="bar" style="width:{tiers[1]*3}px"></div></div>
<div>Tier 2 (Approval): <strong>{tiers[2]}</strong> <div class="bar" style="width:{tiers[2]*3}px"></div></div>
<div>Tier 3 (Human): <strong>{tiers[3]}</strong> <div class="bar" style="width:{tiers[3]*3}px"></div></div>
</div>
<div class="section"><h2>Query Types</h2><table><th>Type</th><th>Count</th>{type_rows}</table></div>
</div>
<div class="section"><h2>Most Read Wiki</h2><table><th>Article</th><th>Reads</th>{wiki_rows}</table></div>
<p style="color:#aaa;font-size:12px;margin-top:16px">Auto-refresh: 30 seconds | Last update: {datetime.now().strftime('%H:%M:%S')}</p>
</body></html>"""
        return HTMLResponse(content=html)

    return app


# =============================================================================
# CLI
# =============================================================================

def load_traces(days=None):
    """Load trace files. days=7 for last 7 days, None for all."""
    import glob as glob_mod
    from datetime import timedelta
    traces = []
    cutoff = None
    if days:
        cutoff = (datetime.now() - timedelta(days=days)).isoformat()
    for f in sorted(glob_mod.glob(f"{CONFIG['trace_dir']}/*.jsonl")):
        with open(f, "r", encoding="utf-8", errors="replace") as fh:
            for line in fh:
                try:
                    entry = json.loads(line)
                    if cutoff and entry.get("timestamp", "") < cutoff:
                        continue
                    traces.append(entry)
                except json.JSONDecodeError:
                    continue
    return traces


def tier_stats():
    """Print tier distribution statistics from trace files."""
    traces = load_traces()
    if not traces:
        print("No trace data found.")
        return

    total = len(traces)
    tiers = {1: 0, 2: 0, 3: 0}
    for t in traces:
        tier = t.get("tier", 1)
        tiers[tier] = tiers.get(tier, 0) + 1

    print(f"Total queries: {total}")
    print(f"  Tier 1 (Autonomous — info/routing): {tiers[1]} ({tiers[1]/total:.0%})")
    print(f"  Tier 2 (Approval — doc generation):  {tiers[2]} ({tiers[2]/total:.0%})")
    print(f"  Tier 3 (Human — out of scope):       {tiers[3]} ({tiers[3]/total:.0%})")

    types = {}
    for t in traces:
        qt = t.get("query_type", "?")
        types[qt] = types.get(qt, 0) + 1
    print(f"\nQuery type distribution:")
    for qt, count in sorted(types.items(), key=lambda x: -x[1]):
        print(f"  {qt}: {count} ({count/total:.0%})")


def tier_promotions():
    """List candidates for promotion from Tier 2 to Tier 1.

    Analyzes success rate of Tier 2 operations from trace data.
    Does not auto-promote — presents data, human decides.
    """
    traces = load_traces()
    if not traces:
        print("No trace data found.")
        return

    # Group Tier 2 operations by tool
    tier2_ops = {}
    for t in traces:
        if t.get("tier", 1) != 2:
            continue
        for tc in t.get("tool_calls", []):
            name = tc["name"]
            sc = t.get("self_check_result", "passed")
            if name not in tier2_ops:
                tier2_ops[name] = {"total": 0, "passed": 0}
            tier2_ops[name]["total"] += 1
            if sc == "passed":
                tier2_ops[name]["passed"] += 1

    print("=" * 50)
    print("  Tier Promotion Analysis")
    print("=" * 50)

    if not tier2_ops:
        print("\nNo Tier 2 operations found.")
        print("=" * 50)
        return

    for tool, stats in sorted(tier2_ops.items(), key=lambda x: -x[1]["total"]):
        total = stats["total"]
        passed = stats["passed"]
        rate = passed / total if total else 0
        status = "CANDIDATE" if rate >= 0.9 and total >= 10 else "WAIT"
        print(f"\n  {tool}:")
        print(f"    Usage: {total}x | Success: {passed}/{total} ({rate:.0%}) | {status}")
        if status == "CANDIDATE":
            print(f"    -> 90%+ success, 10+ uses — can be evaluated for Tier 1 promotion")

    print("\n" + "=" * 50)


def sop_gaps():
    """Detect topics with no corresponding wiki article.

    Collects missing_source and procedure queries without wiki usage from traces,
    reporting which topics need wiki articles to be written.
    """
    traces = load_traces()
    if not traces:
        print("No trace data found.")
        return

    # Collect procedure/deadline queries with missing_source (excluding constraint — those are rejections)
    gaps = []
    for t in traces:
        sc = t.get("self_check_result", "passed")
        qt = t.get("query_type", "?")
        if sc == "missing_source" and qt not in ("constraint",):
            gaps.append(t.get("query", "?"))

    # Procedure/deadline type queries where wiki was not used
    no_wiki = []
    for t in traces:
        qt = t.get("query_type", "?")
        wiki = t.get("wiki_articles_used", [])
        if qt in ("procedure", "deadline") and not wiki:
            no_wiki.append(t.get("query", "?"))

    # Group recurring queries
    from collections import Counter
    gap_counts = Counter(gaps)
    no_wiki_counts = Counter(no_wiki)

    print("=" * 50)
    print("  SOP Gap Report")
    print("=" * 50)

    if gap_counts:
        print(f"\n--- Missing source ({len(gap_counts)} unique topics) ---")
        for query, count in gap_counts.most_common(10):
            print(f"  [{count}x] {query[:80]}")

    if no_wiki_counts:
        print(f"\n--- Answered without wiki ({len(no_wiki_counts)} unique) ---")
        for query, count in no_wiki_counts.most_common(10):
            print(f"  [{count}x] {query[:80]}")

    if not gap_counts and not no_wiki_counts:
        print("\nNo gaps detected — all queries are covered by wiki.")
    print("=" * 50)


def deadline_check(ahead_days=5):
    """List upcoming tax return deadlines.

    Based on monthly dates from beyanname_takvimi.md.
    Reports deadlines within ahead_days.
    """
    from datetime import timedelta

    today = datetime.now()
    target = today + timedelta(days=ahead_days)

    # Monthly tax return calendar (day: description)
    monthly_deadlines = {
        7: "ZUS declaration — small firms (up to 9 employees)",
        10: "ZUS declaration — sole proprietor (JDG, no employees)",
        15: "ZUS declaration — medium/large firms (10+ employees)",
        20: "PIT/CIT advance payment (zaliczka)",
        25: "JPK_V7M / VAT return",
    }

    # Annual tax return calendar (month, day): description
    annual_deadlines = {
        (1, 31): "PIT-11 delivery to employees",
        (2, 28): "PIT-11 submission to tax office",
        (4, 30): "PIT-36 / PIT-37 annual return",
        (6, 30): "CIT-8 annual corporate tax",
    }

    upcoming = []

    # Check monthly deadlines
    for day, desc in monthly_deadlines.items():
        # Bu ayin deadline'i
        try:
            deadline_date = today.replace(day=day)
        except ValueError:
            continue  # Month ends at 28/29 days
        if today.date() <= deadline_date.date() <= target.date():
            days_left = (deadline_date.date() - today.date()).days
            upcoming.append((deadline_date.date(), days_left, desc))

    # Check annual deadlines
    for (month, day), desc in annual_deadlines.items():
        try:
            deadline_date = today.replace(month=month, day=day)
        except ValueError:
            continue
        if today.date() <= deadline_date.date() <= target.date():
            days_left = (deadline_date.date() - today.date()).days
            upcoming.append((deadline_date.date(), days_left, desc))

    upcoming.sort(key=lambda x: x[0])

    print(f"Upcoming deadlines (today: {today.strftime('%Y-%m-%d')}, next {ahead_days} days):")
    if not upcoming:
        print(f"  No deadlines within the next {ahead_days} days.")
    else:
        for date, days_left, desc in upcoming:
            urgency = "⚠" if days_left <= 2 else " "
            print(f"  {urgency} {date} ({days_left} days) — {desc}")
    print(f"\nNote: Deadlines falling on weekends are extended to the next business day.")


def deadline_remind(ahead_days=5):
    """T11: Proactive deadline reminders — generate client-specific drafts.

    1. Detect upcoming deadlines
    2. Filter by firm type from mastersheet (JDG vs Sp. z o.o.)
    3. Generate reminder drafts for each deadline+client pair
    4. Save under outputs/reminders/
    """
    import csv as csv_mod
    from datetime import timedelta

    today = datetime.now()
    target = today + timedelta(days=ahead_days)

    # Deadline types and which firm types they apply to
    deadline_rules = {
        7:  {"desc": "ZUS declaration (small)", "firms": ["SP"]},
        10: {"desc": "ZUS declaration (JDG)", "firms": ["JDG"]},
        15: {"desc": "ZUS declaration (medium/large)", "firms": ["SP"]},
        20: {"desc": "PIT/CIT advance", "firms": ["JDG", "SP"]},
        25: {"desc": "JPK_V7M / VAT", "firms": ["JDG", "SP"]},
    }

    # Find upcoming deadlines
    upcoming = []
    for day, rule in deadline_rules.items():
        try:
            d = today.replace(day=day)
        except ValueError:
            continue
        if today.date() <= d.date() <= target.date():
            days_left = (d.date() - today.date()).days
            upcoming.append((d.date(), days_left, rule["desc"], rule["firms"]))
    upcoming.sort(key=lambda x: x[0])

    if not upcoming:
        print(f"No deadlines within the next {ahead_days} days.")
        return

    # Load clients
    firms = []
    try:
        with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
            for row in csv_mod.DictReader(f):
                name = row.get("Company Name", "").strip()
                typ = row.get("TYP", "").strip().upper()
                ra = row.get("RA", "").strip()
                if name:
                    firm_cat = "JDG" if "JDG" in typ else "SP"
                    firms.append({"name": name, "typ": typ, "cat": firm_cat, "ra": ra})
    except FileNotFoundError:
        print("Mastersheet not found.")
        return

    # Generate reminder drafts
    output_dir = os.path.join(CONFIG["output_dir"], "reminders", today.strftime("%Y-%m-%d"))
    os.makedirs(output_dir, exist_ok=True)

    total_reminders = 0
    for deadline_date, days_left, desc, firm_types in upcoming:
        matching_firms = [f for f in firms if f["cat"] in firm_types]
        if not matching_firms:
            continue

        print(f"\n{deadline_date} ({days_left} days) — {desc}: {len(matching_firms)} firms")

        for firm in matching_firms:
            reminder = (
                f"Sayin {firm['name']} yetkilisi,\n\n"
                f"{desc} icin son tarih {deadline_date} ({days_left} gun kaldi).\n"
                f"Lutfen gerekli belgelerin zamaninda teslim edilmesini saglayiniz.\n\n"
                f"Sorulariniz icin bizimle iletisime gecebilirsiniz.\n\n"
                f"Saygilarimizla,\n"
                f"Support Team"
            )
            total_reminders += 1

        # Batch draft file (per deadline)
        safe_desc = desc.replace("/", "-").replace(" ", "_")[:30]
        filepath = os.path.join(output_dir, f"{deadline_date}_{safe_desc}.txt")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"# {desc} — {deadline_date} ({days_left} days)\n")
            f.write(f"# Total: {len(matching_firms)} firms\n")
            f.write(f"# Created: {today.isoformat()}\n")
            f.write(f"# WARNING: No automatic sending — human approval required\n\n")
            for firm in matching_firms:
                f.write(f"--- {firm['name']} ({firm['typ']}) | Accountant: {firm['ra']} ---\n")
                f.write(f"Sayin {firm['name']} yetkilisi,\n\n")
                f.write(f"{desc} icin son tarih {deadline_date} ({days_left} gun kaldi).\n")
                f.write(f"Lutfen gerekli belgelerin zamaninda teslim edilmesini saglayiniz.\n\n")
                f.write(f"Support Team\n\n")
        print(f"  -> {filepath}")

    print(f"\nTotal: {total_reminders} reminder drafts created.")
    print(f"Location: {output_dir}")
    print(f"IMPORTANT: No automatic sending. Review drafts and send manually.")


def silent_clients(days=30, action=False):
    """List clients not queried in the last N days.
    action=True generates check-in drafts (T12).
    """
    import csv as csv_mod
    from datetime import timedelta

    # T10: Read from client_state first (if exists), otherwise derive from traces
    _load_client_state()
    cutoff = (datetime.now() - timedelta(days=days)).isoformat()

    all_companies = []
    try:
        with open(CONFIG["mastersheet_file"], "r", encoding="utf-8") as f:
            for row in csv_mod.DictReader(f):
                name = row.get("Company Name", "").strip()
                if name:
                    all_companies.append(name)
    except FileNotFoundError:
        print("Mastersheet not found.")
        return

    # Last contact from client state or traces
    silent = []
    for company in all_companies:
        last_ts = None
        # Check client state
        for nip, state in _client_state.items():
            if state.get("company", "").lower() == company.lower():
                last_ts = state.get("last_contact")
                break
        if last_ts is None:
            silent.append((company, "Never queried", None))
        elif last_ts < cutoff:
            days_ago = (datetime.now() - datetime.fromisoformat(last_ts)).days
            silent.append((company, f"Last: {last_ts[:10]}", days_ago))

    print(f"Silent clients (last {days} days):")
    print(f"Total: {len(silent)} / {len(all_companies)}\n")
    for name, status, _ in sorted(silent):
        print(f"  {name} — {status}")

    # T12: Generate check-in drafts with --action flag
    if action and silent:
        output_dir = os.path.join(CONFIG["output_dir"], "checkins", datetime.now().strftime("%Y-%m-%d"))
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, "check_in_taslak.txt")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"# Silent Client Check-in Drafts\n")
            f.write(f"# Date: {datetime.now().strftime('%Y-%m-%d')}\n")
            f.write(f"# WARNING: No automatic sending\n\n")
            for name, status, days_ago in sorted(silent):
                # Adjust tone based on duration
                if days_ago and days_ago >= 90:
                    tone = "ALERT — Notify Kaan Bey (churn risk)"
                    msg = f"Hesabinizla ilgili onemli bir kontrol yapmak istedik. Uzun suredir gorusemedik."
                elif days_ago and days_ago >= 60:
                    tone = "SERIOUS"
                    msg = f"Hesabinizla ilgili kontrol yapmak istedik. Yardimci olabilecegimiz bir konu var mi?"
                else:
                    tone = "NORMAL"
                    msg = f"Merhaba, sizinle uzun suredir gorusemedik. Her sey yolunda mi?"
                f.write(f"--- {name} | {status} | [{tone}] ---\n")
                f.write(f"Sayin {name} yetkilisi,\n\n{msg}\n\nSupport Team\n\n")
        print(f"\nCheck-in drafts: {filepath}")
        # Kaan alert for 90+ day silent clients
        critical = [(n, d) for n, _, d in silent if d and d >= 90]
        if critical:
            print(f"\nCRITICAL — To be reported to Kaan Bey ({len(critical)} firms, 90+ days silent):")
            for n, d in critical:
                print(f"  {n} ({d} days)")


def daily_optimize():
    """End-of-day feedback evaluation and system prompt improvement.

    1. Analyzes today's feedback
    2. Generates QA pairs from negative feedback (run_benchmark --feedback)
    3. Improves system prompt with 1 iteration of hill-climb
    4. Prints daily summary report
    """
    import subprocess

    base = os.path.dirname(os.path.abspath(__file__))
    benchmark_script = os.path.join(base, "run_benchmark.py")

    print("=" * 50)
    print(f"  CS Agent Daily Optimization — {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    print("=" * 50)

    # 1. Read today's feedback
    feedback_path = f"{CONFIG['trace_dir']}/feedback.jsonl"
    today = datetime.now().strftime('%Y-%m-%d')
    pos, neg = 0, 0
    if os.path.exists(feedback_path):
        with open(feedback_path, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    entry = json.loads(line)
                    if entry.get("timestamp", "").startswith(today):
                        if entry.get("rating", 0) > 0:
                            pos += 1
                        else:
                            neg += 1
                except json.JSONDecodeError:
                    continue

    print(f"\nToday's feedback: {pos} positive, {neg} negative")

    if neg == 0 and pos == 0:
        print("No feedback today — skipping optimization.")
        print("=" * 50)
        return

    # 2. Generate QA if there are negative feedback
    if neg > 0:
        print(f"\n[1/3] Generating QA pairs from {neg} negative feedback...")
        result = subprocess.run(
            ["python", benchmark_script, "--feedback"],
            cwd=base, capture_output=True, text=True, timeout=300
        )
        print(result.stdout[-200:] if result.stdout else "  No output")
        if result.returncode != 0:
            print(f"  Error: {result.stderr[-200:]}")
    else:
        print("\n[1/3] No negative feedback — skipping QA generation.")

    # 3. Hill-climb (only if there are negative feedback)
    if neg > 0:
        print("\n[2/3] System prompt improvement (1 iteration)...")
        result = subprocess.run(
            ["python", benchmark_script, "--hill-climb", "1"],
            cwd=base, capture_output=True, text=True, timeout=600
        )
        print(result.stdout[-300:] if result.stdout else "  No output")
        if result.returncode != 0:
            print(f"  Error: {result.stderr[-200:]}")
    else:
        print("\n[2/3] No negatives — skipping hill-climb.")

    # 4. Daily summary
    traces = load_traces(days=1)
    total = len(traces)
    errors = sum(1 for t in traces if t.get("self_check_result") != "passed")

    print(f"\n[3/3] Daily summary:")
    print(f"  Total queries: {total}")
    print(f"  Error rate: {errors}/{total}")
    print(f"  Feedback: +{pos} / -{neg}")
    print("=" * 50)


def weekly_report():
    """Last 7 days CS summary report."""
    traces = load_traces(days=7)
    if not traces:
        print("No trace data found in the last 7 days.")
        return

    total = len(traces)
    tiers = {1: 0, 2: 0, 3: 0}
    types = {}
    checks = {"passed": 0, "missing_source": 0, "ungrounded_claim": 0, "flagged": 0}
    durations = []
    wiki_usage = {}
    routing_targets = {}

    for t in traces:
        tiers[t.get("tier", 1)] = tiers.get(t.get("tier", 1), 0) + 1
        qt = t.get("query_type", "?")
        types[qt] = types.get(qt, 0) + 1
        sc = t.get("self_check_result", "passed")
        checks[sc] = checks.get(sc, 0) + 1
        if t.get("duration_ms"):
            durations.append(t["duration_ms"])
        for art in t.get("wiki_articles_used", []):
            wiki_usage[art] = wiki_usage.get(art, 0) + 1
        rd = t.get("routing_decision")
        if rd:
            routing_targets[rd] = routing_targets.get(rd, 0) + 1

    avg_ms = sum(durations) / len(durations) if durations else 0
    error_count = checks.get("missing_source", 0) + checks.get("ungrounded_claim", 0) + checks.get("flagged", 0)

    print("=" * 50)
    print("  CS Agent — Weekly Report")
    print("=" * 50)
    print(f"\nTotal queries: {total}")
    print(f"Average response time: {avg_ms/1000:.1f}s")
    print(f"Error rate: {error_count}/{total} ({error_count/total:.0%})" if total else "")

    print(f"\n--- Tier Distribution ---")
    print(f"  Tier 1 (Autonomous): {tiers.get(1,0)} ({tiers.get(1,0)/total:.0%})")
    print(f"  Tier 2 (Approval):   {tiers.get(2,0)} ({tiers.get(2,0)/total:.0%})")
    print(f"  Tier 3 (Human):      {tiers.get(3,0)} ({tiers.get(3,0)/total:.0%})")

    print(f"\n--- Query Types ---")
    for qt, count in sorted(types.items(), key=lambda x: -x[1]):
        print(f"  {qt}: {count}")

    if routing_targets:
        print(f"\n--- Routing Targets ---")
        for target, count in sorted(routing_targets.items(), key=lambda x: -x[1]):
            print(f"  {target}: {count}")

    if wiki_usage:
        print(f"\n--- Most Read Wiki ---")
        for art, count in sorted(wiki_usage.items(), key=lambda x: -x[1])[:5]:
            print(f"  {art}: {count}")

    print(f"\n--- Self-Check ---")
    print(f"  Passed: {checks.get('passed',0)}")
    if checks.get("missing_source", 0):
        print(f"  Missing source: {checks['missing_source']}")
    if checks.get("ungrounded_claim", 0):
        print(f"  Ungrounded claim: {checks['ungrounded_claim']}")
    if checks.get("flagged", 0):
        print(f"  Flagged: {checks['flagged']}")
    print("=" * 50)


def trace_trends(days=7):
    """T13: Last N days trend analysis — topic, tool, error trends + anomaly detection."""
    traces = load_traces(days=days)
    if not traces:
        print(f"No trace data found in the last {days} days.")
        return

    # Daily grouping
    daily = {}
    for t in traces:
        day = t.get("timestamp", "")[:10]
        if not day:
            continue
        if day not in daily:
            daily[day] = {"total": 0, "errors": 0, "types": {}, "tools": {}}
        d = daily[day]
        d["total"] += 1
        if t.get("query_type") == "error" or t.get("error_category"):
            d["errors"] += 1
        qt = t.get("query_type", "?")
        d["types"][qt] = d["types"].get(qt, 0) + 1
        for tc in t.get("tool_calls", []):
            tn = tc.get("name", "?")
            d["tools"][tn] = d["tools"].get(tn, 0) + 1

    print(f"{'=' * 55}")
    print(f"  Trend Analysis — Last {days} Days")
    print(f"{'=' * 55}")

    # Daily summary
    print(f"\n--- Daily Queries ---")
    for day in sorted(daily.keys()):
        d = daily[day]
        err_pct = d["errors"] / d["total"] * 100 if d["total"] else 0
        bar = "#" * min(d["total"], 40)
        alert = " !" if err_pct > 10 else ""
        print(f"  {day}: {d['total']:3d} queries | errors: {err_pct:.0f}%{alert} {bar}")

    # Topic trend — totals
    all_types = {}
    for d in daily.values():
        for qt, c in d["types"].items():
            all_types[qt] = all_types.get(qt, 0) + c
    print(f"\n--- Topic Distribution ---")
    for qt, c in sorted(all_types.items(), key=lambda x: -x[1])[:8]:
        print(f"  {qt}: {c}")

    # Tool trend
    all_tools = {}
    for d in daily.values():
        for tn, c in d["tools"].items():
            all_tools[tn] = all_tools.get(tn, 0) + c
    print(f"\n--- Tool Usage ---")
    for tn, c in sorted(all_tools.items(), key=lambda x: -x[1]):
        print(f"  {tn}: {c}")

    # Anomaly detection — compare last day with average
    if len(daily) >= 3:
        days_sorted = sorted(daily.keys())
        last_day = daily[days_sorted[-1]]
        prev_days = [daily[d] for d in days_sorted[:-1]]
        avg_total = sum(d["total"] for d in prev_days) / len(prev_days)
        anomalies = []
        if last_day["total"] > avg_total * 2:
            anomalies.append(f"Query count 2x above average ({last_day['total']} vs avg. {avg_total:.0f})")
        # Topic anomaly
        for qt, c in last_day["types"].items():
            prev_avg = sum(d["types"].get(qt, 0) for d in prev_days) / len(prev_days)
            if prev_avg > 0 and c > prev_avg * 3:
                anomalies.append(f"'{qt}' topic {c}x (avg. {prev_avg:.0f}x) — sudden spike")
        if anomalies:
            print(f"\n--- ANOMALY ---")
            for a in anomalies:
                print(f"  ! {a}")

    print(f"\n{'=' * 55}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "--stats":
        tier_stats()
    elif len(sys.argv) > 1 and sys.argv[1] == "--report":
        weekly_report()
    elif len(sys.argv) > 1 and sys.argv[1] == "--trends":
        days = int(sys.argv[2]) if len(sys.argv) > 2 else 7
        trace_trends(days)
    elif len(sys.argv) > 1 and sys.argv[1] == "--deadlines":
        ahead = int(sys.argv[2]) if len(sys.argv) > 2 else 5
        deadline_check(ahead)
    elif len(sys.argv) > 1 and sys.argv[1] == "--remind":
        ahead = int(sys.argv[2]) if len(sys.argv) > 2 else 5
        deadline_remind(ahead)
    elif len(sys.argv) > 1 and sys.argv[1] == "--silent":
        days = int(sys.argv[2]) if len(sys.argv) > 2 and sys.argv[2].isdigit() else 30
        has_action = "--action" in sys.argv
        silent_clients(days, action=has_action)
    elif len(sys.argv) > 1 and sys.argv[1] == "--gaps":
        sop_gaps()
    elif len(sys.argv) > 1 and sys.argv[1] == "--wiki-rollback":
        # T14: Rollback wiki article to a previous version
        if len(sys.argv) < 3:
            print("Usage: --wiki-rollback <article_name> [timestamp]")
            print("Example: --wiki-rollback onboarding/checklist.md 20260409_143000")
        else:
            article = sys.argv[2]
            versions_dir = os.path.join(CONFIG["wiki_dir"], ".versions", article.replace("/", "_"))
            if not os.path.exists(versions_dir):
                print(f"No version history for '{article}'.")
            elif len(sys.argv) >= 4:
                ts = sys.argv[3]
                backup = os.path.join(versions_dir, f"{ts}.md")
                if os.path.exists(backup):
                    import shutil
                    target = os.path.join(CONFIG["wiki_dir"], article)
                    _wiki_backup(article, target)  # Also backup current version
                    shutil.copy2(backup, target)
                    print(f"Rolled back: {article} -> version {ts}")
                else:
                    print(f"Version not found: {ts}")
                    print(f"Available versions: {sorted(os.listdir(versions_dir))}")
            else:
                versions = sorted(os.listdir(versions_dir))
                print(f"'{article}' versions ({len(versions)}):")
                for v in versions:
                    print(f"  {v}")
    elif len(sys.argv) > 1 and sys.argv[1] == "--promotions":
        tier_promotions()
    elif len(sys.argv) > 1 and sys.argv[1] == "--optimize":
        daily_optimize()
    elif len(sys.argv) > 1 and sys.argv[1] == "--purge-old-traces":
        # S5: GDPR — purge old traces
        days = int(sys.argv[2]) if len(sys.argv) > 2 else 90
        import glob as glob_mod
        from datetime import timedelta
        cutoff = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
        removed = 0
        for f in glob_mod.glob(f"{CONFIG['trace_dir']}/*.jsonl"):
            fname = os.path.basename(f)
            # Dated files: 2026-04-05.jsonl
            if fname[:10] <= cutoff and fname[4] == "-" and fname != "feedback.jsonl":
                os.remove(f)
                removed += 1
                print(f"  Deleted: {fname}")
        print(f"\n{removed} files purged (older than {days} days).")
    elif len(sys.argv) > 1 and sys.argv[1] == "--serve":
        # HTTP API mode
        import uvicorn
        # T3: Load existing sessions when server starts
        _load_sessions()
        # T10: Initialize/load client state
        _init_client_state()
        # T8: Start enabled MCP servers (Gmail, Supabase, etc.)
        _init_mcp_servers()
        app = create_api()
        port = int(sys.argv[2]) if len(sys.argv) > 2 else 8000
        print(f"CS Agent API — http://0.0.0.0:{port}")
        uvicorn.run(app, host="0.0.0.0", port=port)
    elif len(sys.argv) > 1:
        # Single query mode
        query = " ".join(sys.argv[1:])
        session = {"messages": []}
        print(agent_loop(query, session))
    else:
        # REPL mode
        print("CS Agent v1 — Type 'q' to quit")
        session = {"messages": []}
        while True:
            try:
                query = input("\n> ")
            except (EOFError, KeyboardInterrupt):
                break
            if query.lower() in ("exit", "quit", "q"):
                break
            print(agent_loop(query, session))
